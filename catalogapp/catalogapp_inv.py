import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk
from tkinter import scrolledtext
import pandas as pd
import requests
from docx import Document
import webbrowser
import os
from functools import partial


# ----- App Config -----
BASE_URL = "https://secondstate.art"
APP_TITLE = "Art Catalog Uploader"
APP_MIN_W, APP_MIN_H = 1100, 780
CATALOG_API_KEY = "276e19f127f140623e73e6c160bbd8ed"
DEFAULT_CATALOG_PATH = r"I:\Shared drives\SECONDSTATE\THE BOOKS\SUPREME.xlsx"
DEFAULT_CATALOG_SHEET = "Inventory for June 2026"
# DEFAULT_CATALOG_PATH = r"J:\Shared drives\SECONDSTATE\THE BOOKS\Oliver Current Print Inventory - Cataloging Program.xlsx"


def api_headers():
    return {"X-API-KEY": CATALOG_API_KEY}


class ArtCatalogApp:
    def __init__(self, master):
        self.master = master
        master.title(APP_TITLE)
        master.geometry(f"{APP_MIN_W}x{APP_MIN_H}")
        master.minsize(APP_MIN_W, APP_MIN_H)

        # -------- THEME / STYLE --------
        self.style = ttk.Style(master)
        try:
            self.style.theme_use("clam")
        except tk.TclError:
            pass

        ACCENT = "#4E8DF5"
        SUCCESS = "#28a745"
        DANGER = "#d9534f"
        BG = "#F6F7FB"

        master.configure(bg=BG)

        self.style.configure(".", font=("Segoe UI", 11))
        self.style.configure("Header.TLabel", font=("Segoe UI Semibold", 16))
        self.style.configure("Subtle.TLabel", foreground="#5a5f73")
        self.style.configure("TButton", padding=8)
        self.style.map(
            "TButton",
            foreground=[("active", "!disabled", "#000")],
            background=[("active", "!disabled", "#E8ECF7")],
        )
        self.style.configure(
            "Accent.TButton",
            padding=10,
            font=("Segoe UI Semibold", 11),
            foreground="white",
            background=ACCENT,
        )
        self.style.map("Accent.TButton", background=[("active", "!disabled", "#3E78D1")])
        self.style.configure(
            "Success.TButton",
            padding=10,
            font=("Segoe UI Semibold", 11),
            foreground="white",
            background=SUCCESS,
        )
        self.style.map("Success.TButton", background=[("active", "!disabled", "#1f7e37")])
        self.style.configure(
            "Danger.TButton",
            padding=10,
            font=("Segoe UI Semibold", 11),
            foreground="white",
            background=DANGER,
        )
        self.style.map("Danger.TButton", background=[("active", "!disabled", "#b8453f")])

        self.style.configure("Catalog.Treeview", rowheight=28, borderwidth=0, relief="flat")
        self.style.configure("Catalog.Treeview.Heading", font=("Segoe UI Semibold", 11))
        self.style.map(
            "Catalog.Treeview",
            background=[("selected", "#E3EDFF")],
            foreground=[("selected", "#000")],
        )

        self.df = None
        self.current_results = None
        self.open_listings_button = None

        self._build_menubar()

        header = ttk.Frame(master, padding=(16, 10, 16, 10))
        header.pack(fill=tk.X)
        ttk.Label(header, text="🎨 Art Catalog Uploader", style="Header.TLabel").pack(side=tk.LEFT)
        ttk.Label(
            header,
            text="Search, manage, and upload artworks to your website.",
            style="Subtle.TLabel",
        ).pack(side=tk.LEFT, padx=(10, 0))

        paned = ttk.Panedwindow(master, orient=tk.HORIZONTAL)
        paned.pack(fill=tk.BOTH, expand=True, padx=16, pady=(0, 10))

        left = ttk.Frame(paned, padding=12)
        left.configure(style="Card.TFrame")
        paned.add(left, weight=0)

        right = ttk.Frame(paned, padding=0)
        paned.add(right, weight=1)

        ttk.Label(left, text="Using shared catalog database", style="Subtle.TLabel").pack(anchor="w", pady=(0, 10))

        ttk.Label(left, text="Search by Title or Artist").pack(anchor="w", pady=(8, 4))
        self.search_var = tk.StringVar()
        search_row = ttk.Frame(left)
        search_row.pack(fill=tk.X)
        self.search_entry = ttk.Entry(search_row, textvariable=self.search_var)
        self.search_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.search_entry.bind("<Return>", self.search_catalog)
        ttk.Button(search_row, text="Search", command=self.search_catalog, style="Accent.TButton").pack(
            side=tk.LEFT, padx=(8, 0)
        )

        ttk.Label(left, text="Actions").pack(anchor="w", pady=(16, 6))
        ttk.Button(left, text="Export Search → Word", command=self.export_to_word).pack(fill=tk.X, pady=4)
        ttk.Button(left, text="Export Entire Database", command=self.export_entire_database).pack(fill=tk.X, pady=4)
        ttk.Button(left, text="Upload Artwork(s) to Website", command=self.upload_artworks, style="Success.TButton").pack(
            fill=tk.X, pady=(10, 4)
        )
        ttk.Button(left, text="Delete Artwork from Website", command=self.delete_artwork, style="Danger.TButton").pack(
            fill=tk.X, pady=4
        )
        ttk.Button(left, text="Clear ALL Website Listings", command=self.clear_all_website_listings, style="Danger.TButton").pack(
            fill=tk.X, pady=4
        )
        ttk.Separator(left).pack(fill=tk.X, pady=12)
        ttk.Button(left, text="Go to Your Listings", command=lambda: webbrowser.open(f"{BASE_URL}/artworks/")).pack(
            fill=tk.X, pady=4
        )

        ttk.Button(left, text="Display All Inventory", command=self.display_all_inventory, style="Accent.TButton").pack(fill=tk.X, pady=4)

        ttk.Label(left, text="Shortcuts", style="Subtle.TLabel").pack(anchor="w", pady=(16, 6))
        shortcuts = (
            "Ctrl+F: Search\n"
            "Ctrl+E: Export search\n"
            "Ctrl+D: Export database\n"
            "Ctrl+U: Upload results\n"
            "Ctrl+Q: Quit"
        )
        ttk.Label(left, text=shortcuts, justify="left").pack(anchor="w")

        notebook = ttk.Notebook(right)
        notebook.pack(fill=tk.BOTH, expand=True)

        tab_results = ttk.Frame(notebook, padding=10)
        tab_details = ttk.Frame(notebook, padding=10)
        tab_preview = ttk.Frame(notebook, padding=10)
        notebook.add(tab_results, text="Results")
        notebook.add(tab_details, text="Details")
        notebook.add(tab_preview, text="Preview Text")

        tree_frame = ttk.Frame(tab_results)
        tree_frame.pack(fill=tk.BOTH, expand=True)

        cols = ("Artist", "Title", "Year", "Medium", "Catalog #")
        self.tree = ttk.Treeview(tree_frame, columns=cols, show="headings", style="Catalog.Treeview", selectmode="extended")
        for c in cols:
            self.tree.heading(c, text=c, command=partial(self._sort_tree, c, False))
            self.tree.column(c, width=120, stretch=True)
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=vsb.set)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)

        self.tree.bind("<<TreeviewSelect>>", self._on_select_row)

        self.details_text = scrolledtext.ScrolledText(tab_details, wrap=tk.WORD, height=20, font=("Segoe UI", 11))
        self.details_text.pack(fill=tk.BOTH, expand=True)
        self.details_text.configure(state=tk.DISABLED)

        self.text_display = scrolledtext.ScrolledText(tab_preview, wrap=tk.WORD, height=20, font=("Segoe UI", 11))
        self.text_display.pack(fill=tk.BOTH, expand=True)
        self.text_display.configure(state=tk.DISABLED)

        self.status = ttk.Label(master, text="Loading shared catalog database...", anchor="w", padding=(16, 6))
        self.status.pack(fill=tk.X)

        self.master.lift()
        self.master.attributes("-topmost", True)
        self.master.after(600, lambda: self.master.attributes("-topmost", False))

        master.bind_all("<Control-f>", self.search_catalog)
        master.bind_all("<Control-F>", self.search_catalog)
        master.bind_all("<Control-e>", lambda e: self.export_to_word())
        master.bind_all("<Control-E>", lambda e: self.export_to_word())
        master.bind_all("<Control-d>", lambda e: self.export_entire_database())
        master.bind_all("<Control-D>", lambda e: self.export_entire_database())
        master.bind_all("<Control-u>", lambda e: self.upload_artworks())
        master.bind_all("<Control-U>", lambda e: self.upload_artworks())
        master.bind_all("<Control-q>", lambda e: master.quit())
        master.bind_all("<Control-Q>", lambda e: master.quit())

        self.master.after(100, self.load_excel)

    # -------------- UI Helpers --------------
    def _build_menubar(self):
        menubar = tk.Menu(self.master)
        file_menu = tk.Menu(menubar, tearoff=False)
        file_menu.add_command(label="Reload Shared Catalog", command=self.load_excel)
        file_menu.add_separator()
        file_menu.add_command(label="Export Search → Word", accelerator="Ctrl+E", command=self.export_to_word)
        file_menu.add_command(label="Export Entire Database", accelerator="Ctrl+D", command=self.export_entire_database)
        file_menu.add_separator()
        file_menu.add_command(label="Quit", accelerator="Ctrl+Q", command=self.master.quit)
        menubar.add_cascade(label="File", menu=file_menu)

        actions_menu = tk.Menu(menubar, tearoff=False)
        actions_menu.add_command(label="Upload Artwork(s)", accelerator="Ctrl+U", command=self.upload_artworks)
        actions_menu.add_command(label="Delete Artwork from Website", command=self.delete_artwork)
        actions_menu.add_command(label="Clear ALL Website Listings", command=self.clear_all_website_listings)
        menubar.add_cascade(label="Actions", menu=actions_menu)

        help_menu = tk.Menu(menubar, tearoff=False)
        help_menu.add_command(label="Go to Listings", command=lambda: webbrowser.open(f"{BASE_URL}/artworks/"))
        menubar.add_cascade(label="Help", menu=help_menu)

        self.master.config(menu=menubar)

    def _set_status(self, text):
        self.status.config(text=text)

    def _sort_tree(self, col, reverse):
        try:
            data = [(self.tree.set(k, col), k) for k in self.tree.get_children("")]
            if col == "Year":
                def to_int(x):
                    try:
                        return int(x[0])
                    except Exception:
                        return float("inf")
                data.sort(key=to_int, reverse=reverse)
            else:
                data.sort(key=lambda t: (t[0] or "").lower(), reverse=reverse)
            for index, (val, k) in enumerate(data):
                self.tree.move(k, "", index)
            self.tree.heading(col, command=lambda c=col: self._sort_tree(c, not reverse))
        except Exception as e:
            self._set_status(f"Sort error: {e}")

    def display_all_inventory(self):
        if self.df is None:
            messagebox.showwarning("No Catalog", "Open an Excel catalog first.")
            return

        self._clear_results()
        self.current_results = self.df

        for idx, row in self.df.iterrows():
            self.tree.insert(
                "",
                tk.END,
                iid=str(idx),
                values=(
                    self.safe(row.get("Artist", "")),
                    self.safe(row.get("Title", "")),
                    self.safe(row.get("Year", "")),
                    self.safe(row.get("Medium", "")),
                    self.safe(row.get("Catalog Number", "")),
                )
            )

        entries = [self.format_catalog_entry(row) for _, row in self.df.iterrows()]
        self.text_display.configure(state=tk.NORMAL)
        self.text_display.insert(tk.END, "\n\n".join(entries))
        self.text_display.configure(state=tk.DISABLED)

        self._set_status(f"Showing all inventory: {len(self.df)} records.")




    def _on_select_row(self, event=None):
        sel = self.tree.selection()
        if not sel or self.current_results is None:
            return

        idx = int(sel[0])
        row = self.current_results.loc[idx]
        details = self.format_catalog_entry(row)

        self.details_text.configure(state=tk.NORMAL)
        self.details_text.delete("1.0", tk.END)
        self.details_text.insert(tk.END, details)
        self.details_text.configure(state=tk.DISABLED)

    # -------------- Data Loading --------------
    def _load_excel_from_path(self, path):
        preview = pd.read_excel(path, header=None, nrows=30)

        header_row = None
        for i in range(len(preview)):
            row_vals = preview.iloc[i].astype(str).str.strip().tolist()
            if "Title" in row_vals and "Artist" in row_vals:
                header_row = i
                break

        if header_row is None:
            raise ValueError("Could not find a header row containing both 'Title' and 'Artist'.")

        self.df = pd.read_excel(path, header=header_row)
        self.df.columns = [str(c).strip() for c in self.df.columns]
        self._clear_results()
        self._set_status(f"Loaded: {os.path.basename(path)}  —  {len(self.df):,} records")

    def load_excel(self):
        path = DEFAULT_CATALOG_PATH
        try:
            if not os.path.exists(path):
                raise FileNotFoundError(f"Shared catalog file not found:\n{path}")
            self._load_excel_from_path(path)
        except Exception as e:
            messagebox.showerror("Load Error", f"Failed to load Excel file:\n{e}")

    def _clear_results(self):
        for i in self.tree.get_children():
            self.tree.delete(i)
        self.current_results = None
        self.details_text.configure(state=tk.NORMAL)
        self.details_text.delete("1.0", tk.END)
        self.details_text.configure(state=tk.DISABLED)
        self.text_display.configure(state=tk.NORMAL)
        self.text_display.delete("1.0", tk.END)
        self.text_display.configure(state=tk.DISABLED)

    # -------------- Website Listing Helpers --------------
    def _fetch_website_artworks(self):
        response = requests.get(
            f"{BASE_URL}/artworks/?format=json",
            headers=api_headers(),
            timeout=20,
        )
        response.raise_for_status()
        payload = response.json()
        artworks = payload.get("artworks", [])
        if not isinstance(artworks, list):
            raise ValueError("Website returned an unexpected artworks payload.")
        return artworks

    def _clean_name_part(self, value, fallback="UNKNOWN"):
        text = str(value).strip() if value is not None else ""
        return text if text else fallback

    def _artist_last_name(self, artwork):
        artist_candidates = [
            artwork.get("artist"),
            artwork.get("artist_name"),
            artwork.get("artist_display"),
            artwork.get("creator"),
            artwork.get("name"),
        ]
        artist_text = next((self._clean_name_part(v, "") for v in artist_candidates if self._clean_name_part(v, "")), "")
        if not artist_text:
            return "UNKNOWN ARTIST"
        parts = artist_text.replace(",", " ").split()
        return parts[-1].upper() if parts else artist_text.upper()

    def _artwork_title(self, artwork):
        title_candidates = [
            artwork.get("title"),
            artwork.get("name"),
            artwork.get("artwork_title"),
            artwork.get("object_title"),
        ]
        for value in title_candidates:
            title = self._clean_name_part(value, "")
            if title:
                return title
        return "Untitled"

    def _listing_label(self, artwork):
        return f"{self._artist_last_name(artwork)} - {self._artwork_title(artwork)}"

    def _delete_artwork_by_title(self, title):
        response = requests.post(
            f"{BASE_URL}/artworks/delete_artwork/",
            json={"title": title},
            headers=api_headers(),
            timeout=20,
        )
        return response

    def _open_confirm_delete_window(self, artworks_to_delete, parent_window, is_clear_all=False):
        confirm_window = tk.Toplevel(self.master)
        confirm_window.title("Confirm Deletion")
        confirm_window.geometry("760x620")
        confirm_window.minsize(760, 620)
        confirm_window.transient(self.master)
        confirm_window.grab_set()

        frm = ttk.Frame(confirm_window, padding=16)
        frm.pack(fill=tk.BOTH, expand=True)

        heading = "Confirm ALL Website Listings Removal" if is_clear_all else "Confirm Listing Removal"
        ttk.Label(frm, text=heading, style="Header.TLabel").pack(anchor="w")

        warning_text = (
            "This will permanently remove every current artwork listing from SecondState.Art."
            if is_clear_all
            else "These selected listings will be permanently removed from SecondState.Art."
        )
        ttk.Label(frm, text=warning_text, style="Subtle.TLabel", wraplength=640, justify="left").pack(anchor="w", pady=(8, 12))

        count_text = f"{len(artworks_to_delete)} listing(s) selected for deletion:"
        ttk.Label(frm, text=count_text).pack(anchor="w", pady=(0, 6))

        listbox = scrolledtext.ScrolledText(frm, wrap=tk.WORD, height=18, font=("Segoe UI", 11))
        listbox.pack(fill=tk.BOTH, expand=True)
        listbox.insert(tk.END, "\n".join(self._listing_label(art) for art in artworks_to_delete))
        listbox.configure(state=tk.DISABLED)

        progress_label = ttk.Label(frm, text="", style="Subtle.TLabel")
        progress_label.pack(anchor="w", pady=(10, 4))

        btns = ttk.Frame(frm)
        btns.pack(fill=tk.X, pady=(8, 0))

        def run_deletion():
            confirm_btn.config(state=tk.DISABLED)
            cancel_btn.config(state=tk.DISABLED)
            deleted_count = 0
            errors = []

            for index, artwork in enumerate(artworks_to_delete, start=1):
                title = self._artwork_title(artwork)
                label = self._listing_label(artwork)
                progress_label.config(text=f"Removing {index} of {len(artworks_to_delete)}: {label}")
                confirm_window.update_idletasks()
                try:
                    response = self._delete_artwork_by_title(title)
                    if response.status_code == 200:
                        deleted_count += 1
                    else:
                        errors.append(f"{label} — {response.text}")
                except Exception as ex:
                    errors.append(f"{label} — {ex}")

            if parent_window and parent_window.winfo_exists():
                parent_window.destroy()
            confirm_window.destroy()

            if errors:
                preview = "\n".join(errors[:10])
                if len(errors) > 10:
                    preview += f"\n...and {len(errors) - 10} more error(s)."
                messagebox.showwarning(
                    "Deletion Finished with Issues",
                    f"Removed {deleted_count} listing(s).\n\nSome deletions failed:\n{preview}",
                )
                self._set_status(f"Removed {deleted_count} listing(s); {len(errors)} failed.")
            else:
                messagebox.showinfo("Success", f"Removed {deleted_count} listing(s) from SecondState.Art.")
                self._set_status(f"Removed {deleted_count} listing(s) from website.")

        confirm_btn = ttk.Button(
            btns,
            text="Confirm Choices",
            style="Danger.TButton",
            command=run_deletion,
        )
        confirm_btn.pack(side=tk.LEFT)
        cancel_btn = ttk.Button(btns, text="Cancel", command=confirm_window.destroy)
        cancel_btn.pack(side=tk.RIGHT)

        confirm_window.update_idletasks()
        req_w = max(760, min(confirm_window.winfo_reqwidth() + 24, 980))
        req_h = max(620, min(confirm_window.winfo_reqheight() + 24, 820))
        confirm_window.geometry(f"{req_w}x{req_h}")

    # -------------- Core Formatting --------------
    def safe(self, val, fallback=""):
        return str(val).strip() if pd.notna(val) else fallback

    def format_catalog_entry(self, row):
        artist = self.safe(row.get("Artist", "")).upper() or "UNKNOWN ARTIST"
        title = self.safe(row.get("Title", "")) or "Untitled"
        year = self.safe(row.get("Year", ""), "Unknown Year")
        medium = self.safe(row.get("Medium", ""), "Unknown Medium")
        description = self.safe(row.get("Description/Notes", ""))

        height = self.safe(row.get("Height", ""), "?")
        width = self.safe(row.get("Width", ""), "?")
        dimensions_text = f"{height}\" x {width}\""

        catalog_number = self.safe(row.get("Catalog Number", ""), "N/A")
        low_estimate = f"${int(row['Low']):,}" if pd.notna(row.get("Low")) else "N/A"
        high_estimate = f"${int(row['High']):,}" if pd.notna(row.get("High")) else "N/A"
        estimate_text = f"\n\nEstimate: {low_estimate} - {high_estimate}"

        formatted = f"{artist}\n{title}, {year}\n{medium}"
        if description:
            formatted += f"\n{description}"
        formatted += f"\nImage Size: {dimensions_text}"
        if catalog_number:
            formatted += f"\nCatalog #: {catalog_number}"
        formatted += estimate_text

        return formatted

    # -------------- Search --------------
    def search_catalog(self, event=None):
        if self.df is None:
            messagebox.showwarning("No Catalog", "Open an Excel catalog first.")
            return

        query = self.search_var.get().lower().strip()
        if not query:
            messagebox.showwarning("Warning", "Enter a search query!")
            return

        results = self.df[
            self.df["Title"].astype(str).str.lower().str.contains(query, na=False)
            | self.df["Artist"].astype(str).str.lower().str.contains(query, na=False)
        ]

        self._clear_results()

        if results.empty:
            self._set_status("No results found.")
            self.text_display.configure(state=tk.NORMAL)
            self.text_display.insert(tk.END, "No results found.\n")
            self.text_display.configure(state=tk.DISABLED)
            return

        self.current_results = results

        for idx, row in results.iterrows():
            self.tree.insert(
                "",
                tk.END,
                iid=str(idx),
                values=(
                    self.safe(row.get("Artist", "")),
                    self.safe(row.get("Title", "")),
                    self.safe(row.get("Year", "")),
                    self.safe(row.get("Medium", "")),
                    self.safe(row.get("Catalog Number", "")),
                ),
            )

        entries = [self.format_catalog_entry(row) for _, row in results.iterrows()]
        self.text_display.configure(state=tk.NORMAL)
        self.text_display.insert(tk.END, "\n\n".join(entries))
        self.text_display.configure(state=tk.DISABLED)

        self._set_status(f"Found {len(results)} result(s). Select a row to view details.")

    # -------------- Export --------------
    def export_to_word(self):
        text = self.text_display.get("1.0", tk.END).strip()
        if not text or text == "No results found.":
            messagebox.showwarning("Nothing to Export", "Search first, or select text to export.")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not file_path:
            return

        try:
            doc = Document()
            for entry in text.split("\n\n"):
                lines = entry.strip().split("\n")
                if not lines or not lines[0].strip():
                    continue

                artist_p = doc.add_paragraph()
                artist_p.add_run(lines[0]).bold = True

                if len(lines) > 1:
                    title_p = doc.add_paragraph()
                    title_p.add_run(lines[1]).italic = True

                for line in lines[2:]:
                    doc.add_paragraph(line)

                doc.add_paragraph("\n")

            doc.save(file_path)
            messagebox.showinfo("Success", "Exported search results successfully.")
            self._set_status(f"Exported: {os.path.basename(file_path)}")
        except Exception as e:
            messagebox.showerror("Export Error", str(e))

    def export_entire_database(self):
        if self.df is None:
            messagebox.showwarning("No Catalog", "Open an Excel catalog first.")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not file_path:
            return

        try:
            doc = Document()
            for _, row in self.df.iterrows():
                artist = str(row.get("Artist", "")).strip().upper()
                title = str(row.get("Title", "")).strip()
                year = str(row.get("Year", "")).strip()

                p = doc.add_paragraph()
                p.add_run(artist).bold = True
                p.add_run("\n")
                tp = doc.add_paragraph()
                tp.add_run(title).italic = True
                if year:
                    tp.add_run(f", {year}")

                doc.add_paragraph(self.format_catalog_entry(row))
                doc.add_paragraph("\n")

            doc.save(file_path)
            messagebox.showinfo("Success", "Database exported successfully.")
            self._set_status(f"Exported: {os.path.basename(file_path)}")
        except Exception as e:
            messagebox.showerror("Export Error", str(e))

    # -------------- Delete --------------
    def delete_artwork(self):
        delete_window = tk.Toplevel(self.master)
        delete_window.title("Delete Artwork from Website")
        delete_window.geometry("760x620")

        frm = ttk.Frame(delete_window, padding=16)
        frm.pack(fill=tk.BOTH, expand=True)

        ttk.Label(frm, text="Choose Website Listings to Remove", style="Header.TLabel").pack(anchor="w")
        ttk.Label(
            frm,
            text="Check the listings you want to remove, then review your choices on the confirmation screen.",
            style="Subtle.TLabel",
            wraplength=700,
            justify="left",
        ).pack(anchor="w", pady=(8, 12))

        try:
            artwork_list = self._fetch_website_artworks()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to fetch artworks: {e}")
            delete_window.destroy()
            return

        if not artwork_list:
            messagebox.showinfo("No Listings Found", "There are no artwork listings currently on the website.")
            delete_window.destroy()
            return

        canvas = tk.Canvas(frm, highlightthickness=0)
        scrollbar = ttk.Scrollbar(frm, orient="vertical", command=canvas.yview)
        checklist_frame = ttk.Frame(canvas)

        checklist_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all")),
        )
        canvas.create_window((0, 0), window=checklist_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        items = []
        for artwork in artwork_list:
            var = tk.BooleanVar(value=False)
            label = self._listing_label(artwork)
            cb = ttk.Checkbutton(checklist_frame, text=label, variable=var)
            cb.pack(anchor="w", fill=tk.X, pady=2)
            items.append((var, artwork))

        controls = ttk.Frame(frm)
        controls.pack(fill=tk.X, pady=(12, 0))

        info_label = ttk.Label(controls, text=f"{len(artwork_list)} website listing(s) loaded.", style="Subtle.TLabel")
        info_label.pack(side=tk.LEFT)

        def select_all():
            for var, _ in items:
                var.set(True)

        def clear_selection():
            for var, _ in items:
                var.set(False)

        def review_selected():
            selected = [artwork for var, artwork in items if var.get()]
            if not selected:
                messagebox.showwarning("No Selection", "Check at least one listing to remove.")
                return
            self._open_confirm_delete_window(selected, delete_window, is_clear_all=False)

        button_row = ttk.Frame(frm)
        button_row.pack(fill=tk.X, pady=(10, 0))
        ttk.Button(button_row, text="Select All", command=select_all).pack(side=tk.LEFT)
        ttk.Button(button_row, text="Clear Selection", command=clear_selection).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Button(button_row, text="Remove from SecondState.Art", style="Danger.TButton", command=review_selected).pack(side=tk.LEFT, padx=(16, 0))
        ttk.Button(button_row, text="Close", command=delete_window.destroy).pack(side=tk.RIGHT)

    def clear_all_website_listings(self):
        try:
            artwork_list = self._fetch_website_artworks()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to fetch artworks: {e}")
            return

        if not artwork_list:
            messagebox.showinfo("No Listings Found", "There are no artwork listings currently on the website.")
            return

        first_confirm = messagebox.askyesno(
            "Clear ALL Website Listings",
            f"This will remove all {len(artwork_list)} current artwork listing(s) from SecondState.Art.\n\nDo you want to continue to the confirmation screen?",
            icon="warning",
        )
        if not first_confirm:
            return

        self._open_confirm_delete_window(artwork_list, parent_window=None, is_clear_all=True)

    # -------------- Upload --------------
    def upload_artworks(self):
        if self.current_results is None or self.current_results.empty:
            messagebox.showwarning("Warning", "No artworks available. Search or display inventory first.")

            return

        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showwarning("No Selection", "Select one or more artworks from the results list first.")
            return

        if len(selected_items) > 4:
            messagebox.showwarning("Warning", "Can only upload a max of 4 selected items at a time.")
            return

        selected_rows = self.current_results.loc[[int(iid) for iid in selected_items]]

        for _, row in selected_rows.iterrows():
            title = row.get("Title", "")
            price = simpledialog.askstring("Enter Sale Price", f"Enter sale price for ‘{title}’:")
            if not price:
                continue

            # First: select a single cover photo
            images = []

            while True:
                image_path = filedialog.askopenfilename(
                    title=f"Select Image #{len(images) + 1} for ‘{title}’",
                    filetypes=[("Image Files", "*.jpg *.jpeg *.png")]
                )

                if not image_path:
                    if not images:
                        messagebox.showwarning("No Image Selected", "You need to select at least one image.")
                    break

                images.append(image_path)

                more_photos = messagebox.askyesno(
                    "Upload More Photos?",
                    f"{len(images)} image(s) selected for ‘{title}’.\n\nDo you want to add another photo?"
                )

                if not more_photos:
                    break

            if not images:
                continue

            # Detect sheet size flag "S" in Height/Width
            raw_height = str(row.get('Height', '')).strip()
            raw_width = str(row.get('Width', '')).strip()

            if raw_height.startswith('S'):
                height_clean = raw_height.replace('S', '').strip()
                width_clean = raw_width.replace('S', '').strip()
                dimensions_text = ""
                sheet_size = f'{height_clean}" x {width_clean}"'
            else:
                dimensions_text = f'{raw_height}" x {raw_width}"'
                sheet_size = ""

            data = {
                'artist': row.get('Artist', ''),
                'title': title,
                'year': row.get('Year', ''),
                'medium': row.get('Medium', ''),
                'description': row.get('Description/Notes', ''),
                'dimensions_text': dimensions_text,
                'sheet_size': sheet_size,
                'catalog_number': row.get('Catalog Number', ''),
                'price': price,
            }

            files_dict = {}
            try:
                for i, img in enumerate(images):
                    mime = "image/jpeg"
                    if img.lower().endswith(".png"):
                        mime = "image/png"
                    with open(img, 'rb') as f:
                        files_dict[f'image_{i}'] = (os.path.basename(img), f.read(), mime)

                response = requests.post(
                    f"{BASE_URL}/artworks/upload_artwork/",
                    data=data,
                    files=files_dict,
                    headers=api_headers(),
                    timeout=60
                )

                if response.status_code == 201:
                    messagebox.showinfo("Success", f"‘{title}’ uploaded successfully!")
                    self._set_status(f"Uploaded “{title}”.")
                else:
                    messagebox.showerror("Error", f"Upload failed: {response.text}")

            except Exception as e:
                messagebox.showerror("Error", f"Upload failed: {e}")


# -------------- Main --------------
def main():
    root = tk.Tk()
    app = ArtCatalogApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
