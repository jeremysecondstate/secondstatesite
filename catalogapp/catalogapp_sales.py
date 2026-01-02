import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from tkinter import scrolledtext
import pandas as pd
import requests
from docx import Document
import webbrowser
import os
from functools import partial


# ----- App Config -----
BASE_URL = "https://secondstate.art"
APP_TITLE = "Pieces Sold Uploader"
APP_MIN_W, APP_MIN_H = 1100, 780
CATALOG_API_KEY = "276e19f127f140623e73e6c160bbd8ed"


def api_headers():
    return {"X-API-KEY": CATALOG_API_KEY}


class ArtCatalogApp:
    def __init__(self, master):
        self.master = master
        master.title(APP_TITLE)
        master.geometry(f"{APP_MIN_W}x{APP_MIN_H}")
        master.minsize(APP_MIN_W, APP_MIN_H)

        # -------- THEME / STYLE --------
        self.style = ttk.Style(master)
        try:
            self.style.theme_use("clam")
        except tk.TclError:
            pass

        ACCENT = "#4E8DF5"
        SUCCESS = "#28a745"
        DANGER = "#d9534f"
        BG = "#F6F7FB"

        master.configure(bg=BG)

        self.style.configure(".", font=("Segoe UI", 11))
        self.style.configure("Header.TLabel", font=("Segoe UI Semibold", 16))
        self.style.configure("Subtle.TLabel", foreground="#5a5f73")
        self.style.configure("TButton", padding=8)
        self.style.map(
            "TButton",
            foreground=[("active", "!disabled", "#000")],
            background=[("active", "!disabled", "#E8ECF7")],
        )
        self.style.configure(
            "Accent.TButton",
            padding=10,
            font=("Segoe UI Semibold", 11),
            foreground="white",
            background=ACCENT,
        )
        self.style.map("Accent.TButton", background=[("active", "!disabled", "#3E78D1")])

        self.style.configure(
            "Success.TButton",
            padding=10,
            font=("Segoe UI Semibold", 11),
            foreground="white",
            background=SUCCESS,
        )
        self.style.map("Success.TButton", background=[("active", "!disabled", "#1f7e37")])

        self.style.configure(
            "Danger.TButton",
            padding=10,
            font=("Segoe UI Semibold", 11),
            foreground="white",
            background=DANGER,
        )
        self.style.map("Danger.TButton", background=[("active", "!disabled", "#b8453f")])

        self.style.configure("Catalog.Treeview", rowheight=28, borderwidth=0, relief="flat")
        self.style.configure("Catalog.Treeview.Heading", font=("Segoe UI Semibold", 11))
        self.style.map(
            "Catalog.Treeview",
            background=[("selected", "#E3EDFF")],
            foreground=[("selected", "#000")],
        )

        # ------- State -------
        self.df = None
        self.current_results = None

        # ------- MENUBAR -------
        self._build_menubar()

        # ------- HEADER -------
        header = ttk.Frame(master, padding=(16, 10, 16, 10))
        header.pack(fill=tk.X)
        ttk.Label(header, text="🧾 Pieces Sold Uploader", style="Header.TLabel").pack(side=tk.LEFT)
        ttk.Label(
            header,
            text="Search a sales sheet and upload sold pieces to /pieces-sold/.",
            style="Subtle.TLabel",
        ).pack(side=tk.LEFT, padx=(10, 0))

        # ------- MAIN LAYOUT: Paned Window ------
        paned = ttk.Panedwindow(master, orient=tk.HORIZONTAL)
        paned.pack(fill=tk.BOTH, expand=True, padx=16, pady=(0, 10))

        left = ttk.Frame(paned, padding=12)
        left.configure(style="Card.TFrame")
        paned.add(left, weight=0)

        right = ttk.Frame(paned, padding=0)
        paned.add(right, weight=1)

        # ------- LEFT / CONTROL CARD -------
        loader_row = ttk.Frame(left)
        loader_row.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(loader_row, text="Sales Sheet:", style="Subtle.TLabel").pack(side=tk.LEFT)
        ttk.Button(
            loader_row,
            text="Open Excel…",
            command=self.load_excel,
            style="Accent.TButton",
        ).pack(side=tk.RIGHT)

        ttk.Label(left, text="Search by Title/Name or Artist").pack(anchor="w", pady=(8, 4))
        self.search_var = tk.StringVar()
        search_row = ttk.Frame(left)
        search_row.pack(fill=tk.X)
        self.search_entry = ttk.Entry(search_row, textvariable=self.search_var)
        self.search_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.search_entry.bind("<Return>", self.search_catalog)
        ttk.Button(
            search_row,
            text="Search",
            command=self.search_catalog,
            style="Accent.TButton",
        ).pack(side=tk.LEFT, padx=(8, 0))

        ttk.Label(left, text="Actions").pack(anchor="w", pady=(16, 6))
        ttk.Button(left, text="Export Search → Word", command=self.export_to_word).pack(fill=tk.X, pady=4)
        ttk.Button(left, text="Export Entire Sheet → Word", command=self.export_entire_database).pack(fill=tk.X, pady=4)

        ttk.Button(
            left,
            text="Upload Sold Piece(s) to Website",
            command=self.upload_sold_pieces,
            style="Success.TButton",
        ).pack(fill=tk.X, pady=(10, 4))

        ttk.Button(
            left,
            text="Delete Sold Piece from Website",
            command=self.delete_sold_piece,
            style="Danger.TButton",
        ).pack(fill=tk.X, pady=4)

        ttk.Separator(left).pack(fill=tk.X, pady=12)
        ttk.Button(
            left,
            text="Go to Pieces Sold Page",
            command=lambda: webbrowser.open(f"{BASE_URL}/pieces-sold/"),
        ).pack(fill=tk.X, pady=4)

        ttk.Label(left, text="Shortcuts", style="Subtle.TLabel").pack(anchor="w", pady=(16, 6))
        shortcuts = (
            "Ctrl+F: Search\n"
            "Ctrl+E: Export search\n"
            "Ctrl+D: Export entire sheet\n"
            "Ctrl+U: Upload search results\n"
            "Ctrl+Q: Quit"
        )
        ttk.Label(left, text=shortcuts, justify="left").pack(anchor="w")

        # ------- RIGHT / RESULTS NOTEBOOK -------
        notebook = ttk.Notebook(right)
        notebook.pack(fill=tk.BOTH, expand=True)

        tab_results = ttk.Frame(notebook, padding=10)
        tab_details = ttk.Frame(notebook, padding=10)
        tab_preview = ttk.Frame(notebook, padding=10)
        notebook.add(tab_results, text="Results")
        notebook.add(tab_details, text="Details")
        notebook.add(tab_preview, text="Preview Text")

        tree_frame = ttk.Frame(tab_results)
        tree_frame.pack(fill=tk.BOTH, expand=True)

        # Columns geared to SUPREME SALES-like sheet
        cols = ("Artist", "Name", "Date", "Net Sale Price $", "Sale Location")
        self.tree = ttk.Treeview(tree_frame, columns=cols, show="headings", style="Catalog.Treeview")
        for c in cols:
            self.tree.heading(c, text=c, command=partial(self._sort_tree, c, False))
            self.tree.column(c, width=160, stretch=True)
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=vsb.set)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)

        self.tree.bind("<<TreeviewSelect>>", self._on_select_row)

        self.details_text = scrolledtext.ScrolledText(tab_details, wrap=tk.WORD, height=20, font=("Segoe UI", 11))
        self.details_text.pack(fill=tk.BOTH, expand=True)
        self.details_text.configure(state=tk.DISABLED)

        self.text_display = scrolledtext.ScrolledText(tab_preview, wrap=tk.WORD, height=20, font=("Segoe UI", 11))
        self.text_display.pack(fill=tk.BOTH, expand=True)
        self.text_display.configure(state=tk.DISABLED)

        self.status = ttk.Label(master, text="Open a sales Excel sheet to get started.", anchor="w", padding=(16, 6))
        self.status.pack(fill=tk.X)

        self.master.lift()
        self.master.attributes("-topmost", True)
        self.master.after(600, lambda: self.master.attributes("-topmost", False))

        # ------- KEYBOARD SHORTCUTS -------
        master.bind_all("<Control-f>", self.search_catalog)
        master.bind_all("<Control-F>", self.search_catalog)
        master.bind_all("<Control-e>", lambda e: self.export_to_word())
        master.bind_all("<Control-E>", lambda e: self.export_to_word())
        master.bind_all("<Control-d>", lambda e: self.export_entire_database())
        master.bind_all("<Control-D>", lambda e: self.export_entire_database())
        master.bind_all("<Control-u>", lambda e: self.upload_sold_pieces())
        master.bind_all("<Control-U>", lambda e: self.upload_sold_pieces())
        master.bind_all("<Control-q>", lambda e: master.quit())
        master.bind_all("<Control-Q>", lambda e: master.quit())

    # -------------- UI Helpers --------------
    def _build_menubar(self):
        menubar = tk.Menu(self.master)

        file_menu = tk.Menu(menubar, tearoff=False)
        file_menu.add_command(label="Open Excel…", accelerator="Ctrl+O", command=self.load_excel)
        file_menu.add_separator()
        file_menu.add_command(label="Export Search → Word", accelerator="Ctrl+E", command=self.export_to_word)
        file_menu.add_command(label="Export Entire Sheet → Word", accelerator="Ctrl+D", command=self.export_entire_database)
        file_menu.add_separator()
        file_menu.add_command(label="Quit", accelerator="Ctrl+Q", command=self.master.quit)
        menubar.add_cascade(label="File", menu=file_menu)

        actions_menu = tk.Menu(menubar, tearoff=False)
        actions_menu.add_command(label="Upload Sold Piece(s)", accelerator="Ctrl+U", command=self.upload_sold_pieces)
        actions_menu.add_command(label="Delete Sold Piece from Website", command=self.delete_sold_piece)
        menubar.add_cascade(label="Actions", menu=actions_menu)

        help_menu = tk.Menu(menubar, tearoff=False)
        help_menu.add_command(label="Go to Pieces Sold Page", command=lambda: webbrowser.open(f"{BASE_URL}/pieces-sold/"))
        menubar.add_cascade(label="Help", menu=help_menu)

        self.master.config(menu=menubar)

        self.master.bind_all("<Control-o>", lambda e: self.load_excel())
        self.master.bind_all("<Control-O>", lambda e: self.load_excel())

    def _set_status(self, text):
        self.status.config(text=text)

    def _sort_tree(self, col, reverse):
        try:
            data = [(self.tree.set(k, col), k) for k in self.tree.get_children("")]
            data.sort(key=lambda t: (t[0] or "").lower(), reverse=reverse)
            for index, (_, k) in enumerate(data):
                self.tree.move(k, "", index)
            self.tree.heading(col, command=lambda c=col: self._sort_tree(c, not reverse))
        except Exception as e:
            self._set_status(f"Sort error: {e}")

    def _on_select_row(self, event=None):
        sel = self.tree.selection()
        if not sel or self.current_results is None:
            return

        idx = int(sel[0])
        row = self.current_results.loc[idx]
        details = self.format_catalog_entry(row)

        self.details_text.configure(state=tk.NORMAL)
        self.details_text.delete("1.0", tk.END)
        self.details_text.insert(tk.END, details)
        self.details_text.configure(state=tk.DISABLED)

    # -------------- Data Loading --------------
    def load_excel(self):
        path = filedialog.askopenfilename(
            title="Select Sales Excel",
            filetypes=[("Excel files", "*.xlsx *.xls")]
        )
        if not path:
            return

        try:
            preview = pd.read_excel(path, header=None, nrows=30)

            header_row = None
            for i in range(len(preview)):
                row_vals = preview.iloc[i].astype(str).str.strip().tolist()
                if "Name" in row_vals and "Artist" in row_vals and "Date" in row_vals:
                    header_row = i
                    break

            if header_row is None:
                raise ValueError("Could not find a header row containing 'Date', 'Artist', and 'Name'.")

            self.df = pd.read_excel(path, header=header_row)
            self.df.columns = [str(c).strip() for c in self.df.columns]

            # optional: drop Unnamed first col
            if len(self.df.columns) and str(self.df.columns[0]).startswith("Unnamed"):
                self.df = self.df.drop(columns=[self.df.columns[0]])

            try:
                with open(os.path.join(os.path.dirname(__file__), ".last_catalog_path.txt"), "w", encoding="utf-8") as f:
                    f.write(path)
            except Exception:
                pass

            self._clear_results()
            self._set_status(f"Loaded: {os.path.basename(path)}  —  {len(self.df):,} records")

        except Exception as e:
            messagebox.showerror("Load Error", f"Failed to load Excel file:\n{e}")

    def _clear_results(self):
        for i in self.tree.get_children():
            self.tree.delete(i)
        self.current_results = None
        self.details_text.configure(state=tk.NORMAL)
        self.details_text.delete("1.0", tk.END)
        self.details_text.configure(state=tk.DISABLED)
        self.text_display.configure(state=tk.NORMAL)
        self.text_display.delete("1.0", tk.END)
        self.text_display.configure(state=tk.DISABLED)

    # -------------- Core Formatting --------------
    def safe(self, val, fallback=""):
        return str(val).strip() if pd.notna(val) else fallback

    def _clean_money(self, val):
        """Turn '$1,234.00' or '1234' or nan into a plain string or ''."""
        if val is None or (isinstance(val, float) and pd.isna(val)):
            return ""
        s = str(val).strip()
        if s.lower() in ("nan", "none", ""):
            return ""
        s = s.replace("$", "").replace(",", "").strip()
        return s

    def format_catalog_entry(self, row):
        date = str(row.get("Date", ""))[:10] if pd.notna(row.get("Date")) else ""
        artist = self.safe(row.get("Artist", "")).upper() or "UNKNOWN ARTIST"
        title = self.safe(row.get("Name", "")) or "Untitled"
        sale_location = self.safe(row.get("Sale Location", ""), "Unknown Sale Location")
        sold_hammer_price = self.safe(row.get("Sold Hammer Price $", ""))
        link_to_sale = self.safe(row.get("Link to Sale", ""))

        lines = [
            artist,
            f"{title}",
            f"Date Sold: {date}",
            f"Sold Hammer Price: {sold_hammer_price}",
            f"Sale Location: {sale_location}",
            f"Link to Sale: {link_to_sale}" if link_to_sale else "",
        ]
        return "\n".join([l for l in lines if l])

    # -------------- Search --------------
    def search_catalog(self, event=None):
        if self.df is None:
            messagebox.showwarning("No Sales Sheet", "Open an Excel sales sheet first.")
            return

        query = self.search_var.get().lower().strip()
        if not query:
            messagebox.showwarning("Warning", "Enter a search query!")
            return

        # "Name" is title in SUPREME SALES
        name_col = "Name" if "Name" in self.df.columns else "Title"

        results = self.df[
            self.df[name_col].astype(str).str.lower().str.contains(query, na=False) |
            self.df["Artist"].astype(str).str.lower().str.contains(query, na=False)
        ]

        self._clear_results()

        if results.empty:
            self._set_status("No results found.")
            self.text_display.configure(state=tk.NORMAL)
            self.text_display.insert(tk.END, "No results found.\n")
            self.text_display.configure(state=tk.DISABLED)
            return

        self.current_results = results

        for idx, row in results.iterrows():
            date = str(row.get("Date", ""))[:10] if pd.notna(row.get("Date")) else ""
            self.tree.insert(
                "",
                tk.END,
                iid=str(idx),
                values=(
                    self.safe(row.get("Artist", "")),
                    self.safe(row.get(name_col, "")),
                    date,
                    self.safe(row.get("Sold Hammer Price $", "")),
                    self.safe(row.get("Sale Location", "")),
                ),
            )

        entries = [self.format_catalog_entry(row) for _, row in results.iterrows()]
        self.text_display.configure(state=tk.NORMAL)
        self.text_display.insert(tk.END, "\n\n".join(entries))
        self.text_display.configure(state=tk.DISABLED)

        self._set_status(f"Found {len(results)} result(s). Select a row to view details.")

    # -------------- Export --------------
    def export_to_word(self):
        text = self.text_display.get("1.0", tk.END).strip()
        if not text or text == "No results found.":
            messagebox.showwarning("Nothing to Export", "Search first, or select text to export.")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not file_path:
            return

        try:
            doc = Document()
            for entry in text.split("\n\n"):
                lines = entry.strip().split("\n")
                if not lines or not lines[0].strip():
                    continue

                p = doc.add_paragraph()
                p.add_run(lines[0]).bold = True
                for line in lines[1:]:
                    doc.add_paragraph(line)
                doc.add_paragraph("\n")

            doc.save(file_path)
            messagebox.showinfo("Success", "Exported search results successfully.")
            self._set_status(f"Exported: {os.path.basename(file_path)}")
        except Exception as e:
            messagebox.showerror("Export Error", str(e))

    def export_entire_database(self):
        if self.df is None:
            messagebox.showwarning("No Sales Sheet", "Open an Excel sales sheet first.")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not file_path:
            return

        try:
            doc = Document()
            for _, row in self.df.iterrows():
                entry = self.format_catalog_entry(row)
                lines = entry.split("\n")
                if not lines:
                    continue
                p = doc.add_paragraph()
                p.add_run(lines[0]).bold = True
                for line in lines[1:]:
                    doc.add_paragraph(line)
                doc.add_paragraph("\n")

            doc.save(file_path)
            messagebox.showinfo("Success", "Sheet exported successfully.")
            self._set_status(f"Exported: {os.path.basename(file_path)}")
        except Exception as e:
            messagebox.showerror("Export Error", str(e))

    # -------------- Upload (Pieces Sold) --------------
    def upload_sold_pieces(self):
        if self.current_results is None or self.current_results.empty:
            messagebox.showwarning("Warning", "No sold pieces to upload. Search first!")
            return

        if len(self.current_results) > 10:
            messagebox.showwarning("Warning", "Upload max 10 at a time. Narrow your search.")
            return

        ok = 0
        failed = 0

        name_col = "Name" if "Name" in self.current_results.columns else "Title"

        for _, row in self.current_results.iterrows():
            title = self.safe(row.get(name_col, ""))
            artist = self.safe(row.get("Artist", ""))

            if not title:
                continue

            # 1) Pick images (at least 1)
            images = []
            while True:
                image_path = filedialog.askopenfilename(
                    title=f"Select Image for ‘{title}’",
                    filetypes=[("Image Files", "*.jpg *.jpeg *.png")]
                )
                if image_path:
                    images.append(image_path)
                    more = messagebox.askyesno("Upload More?", "Upload another image for this sold piece?")
                    if not more:
                        break
                else:
                    break

            if not images:
                messagebox.showwarning("Missing Images", f"Skipping ‘{title}’ — you must select at least 1 image.")
                failed += 1
                continue

            # 2) Payload (keep your cleaned money + field mapping)
            payload = {
                "date": str(row.get("Date", ""))[:10] if pd.notna(row.get("Date")) else "",
                "artist": artist,
                "title": title,
                "sale_location": self.safe(row.get("Sale Location", "")),
                "sold_hammer_price": self._clean_money(row.get("Sold Hammer Price $", "")),
                "link_to_sale": self.safe(row.get("Link to Sale", "")),
            }

            # 3) Build files dict
            files = {}
            try:
                for i, img in enumerate(images):
                    mime = "image/png" if img.lower().endswith(".png") else "image/jpeg"
                    with open(img, "rb") as f:
                        files[f"image_{i}"] = (os.path.basename(img), f.read(), mime)

                r = requests.post(
                    f"{BASE_URL}/pieces-sold/upload/",
                    data=payload,
                    files=files,
                    headers=api_headers(),
                    timeout=60,
                )

                print("UPLOAD RESPONSE:", r.status_code, r.text[:300])

                if r.status_code == 201:
                    ok += 1
                else:
                    failed += 1
                    messagebox.showerror(
                        "Upload failed",
                        f"{artist} — {title}\n\n{r.status_code}\n{r.text}",
                    )
            except Exception as e:
                failed += 1
                messagebox.showerror("Upload request failed", str(e))

        self._set_status(f"Upload complete. Success: {ok}, Failed: {failed}")

        if ok:
            messagebox.showinfo("Done", f"Uploaded {ok} sold piece(s).")
        elif failed:
            messagebox.showwarning("Upload finished",
                                   f"No uploads succeeded.\nFailed: {failed}\n\nCheck the status bar + any error popups.")
        else:
            messagebox.showinfo("Nothing uploaded", "No items were uploaded (no title or no images were selected).")

    # -------------- Delete (Pieces Sold) --------------
    def delete_sold_piece(self):
        try:
            r = requests.get(f"{BASE_URL}/pieces-sold/?format=json", headers=api_headers(), timeout=15)
            r.raise_for_status()
            sold_list = r.json().get("sold", [])
        except Exception as e:
            messagebox.showerror("Error", f"Failed to fetch sold pieces: {e}")
            return

        if not sold_list:
            messagebox.showinfo("Nothing to delete", "No sold pieces found on the site.")
            return

        win = tk.Toplevel(self.master)
        win.title("Delete Sold Piece")
        win.geometry("620x260")
        frm = ttk.Frame(win, padding=16)
        frm.pack(fill=tk.BOTH, expand=True)

        ttk.Label(frm, text="Select Sold Piece to Delete:", style="Header.TLabel").pack(anchor="w")

        display = []
        id_by_display = {}
        for s in sold_list:
            txt = f"{s.get('artist','')} — {s.get('title','')} ({str(s.get('date',''))})"
            display.append(txt)
            id_by_display[txt] = s.get("id")

        sel_var = tk.StringVar()
        dropdown = ttk.Combobox(frm, textvariable=sel_var, values=display, state="readonly")
        dropdown.pack(fill=tk.X, pady=(10, 6))

        def confirm():
            key = sel_var.get()
            sold_id = id_by_display.get(key)
            if not sold_id:
                return
            try:
                rr = requests.post(
                    f"{BASE_URL}/pieces-sold/delete/",
                    json={"id": sold_id},
                    headers=api_headers(),
                    timeout=20,
                )
                if rr.status_code == 200:
                    messagebox.showinfo("Success", "Deleted sold piece.")
                    win.destroy()
                else:
                    messagebox.showerror("Error", rr.text)
            except Exception as ex:
                messagebox.showerror("Error", str(ex))

        btns = ttk.Frame(frm)
        btns.pack(fill=tk.X, pady=(12, 0))
        ttk.Button(btns, text="Delete", style="Danger.TButton", command=confirm).pack(side=tk.LEFT)
        ttk.Button(btns, text="Close", command=win.destroy).pack(side=tk.RIGHT)


def main():
    root = tk.Tk()
    app = ArtCatalogApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
