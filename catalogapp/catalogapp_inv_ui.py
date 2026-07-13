import os
import threading
import webbrowser
from datetime import datetime
from functools import partial

import pandas as pd
import requests
import tkinter as tk
from docx import Document
from tkinter import filedialog, messagebox, simpledialog, ttk
from tkinter import scrolledtext

BASE_URL = os.environ.get("SECONDSTATE_BASE_URL", "https://secondstate.art").rstrip("/")
APP_TITLE = "Art Catalog Uploader"
APP_MIN_W, APP_MIN_H = 1180, 820
CATALOG_API_KEY = os.environ.get("CATALOG_API_KEY", "276e19f127f140623e73e6c160bbd8ed")
DEFAULT_CATALOG_PATH = r"I:\Shared drives\SECONDSTATE\THE BOOKS\SUPREME.xlsx"
DEFAULT_CATALOG_SHEET = "Inventory for July 2026"


def api_headers(extra=None):
    headers = {"X-API-KEY": CATALOG_API_KEY}
    if extra:
        headers.update(extra)
    return headers


class ArtCatalogApp:
    EDIT_FIELDS = [
        ("artist", "Artist"), ("title", "Title"), ("year", "Year"), ("medium", "Medium"),
        ("dimensions_text", "Image size"), ("sheet_size", "Sheet size"),
        ("catalog_number", "Literature"), ("price", "Price"),
    ]

    def __init__(self, master):
        self.master = master
        self.df = None
        self.current_results = None
        master.title(APP_TITLE)
        master.geometry(f"{APP_MIN_W}x{APP_MIN_H}")
        master.minsize(APP_MIN_W, APP_MIN_H)
        self._style()
        self._build_ui()
        master.bind_all("<Control-f>", self.search_catalog)
        master.bind_all("<Control-e>", lambda _e: self.export_to_word())
        master.bind_all("<Control-d>", lambda _e: self.export_entire_database())
        master.bind_all("<Control-u>", lambda _e: self.upload_artworks())
        master.bind_all("<Control-q>", lambda _e: master.quit())
        master.after(100, self.load_excel)

    def _style(self):
        self.style = ttk.Style(self.master)
        try:
            self.style.theme_use("clam")
        except tk.TclError:
            pass
        self.style.configure(".", font=("Segoe UI", 11))
        self.style.configure("Header.TLabel", font=("Segoe UI Semibold", 16))
        self.style.configure("Subtle.TLabel", foreground="#5a5f73")
        self.style.configure("Accent.TButton", padding=9, font=("Segoe UI Semibold", 11), foreground="white", background="#4E8DF5")
        self.style.configure("Success.TButton", padding=9, font=("Segoe UI Semibold", 11), foreground="white", background="#28a745")
        self.style.configure("Danger.TButton", padding=9, font=("Segoe UI Semibold", 11), foreground="white", background="#d9534f")
        self.style.configure("Catalog.Treeview", rowheight=28)
        self.style.configure("Catalog.Treeview.Heading", font=("Segoe UI Semibold", 11))

    def _build_ui(self):
        menubar = tk.Menu(self.master)
        actions = tk.Menu(menubar, tearoff=False)
        actions.add_command(label="Upload Artwork(s)", command=self.upload_artworks)
        actions.add_command(label="Generate Description for Selected", command=self.generate_description_for_selected)
        actions.add_command(label="Edit Website Listings", command=self.edit_website_listings)
        actions.add_command(label="Delete Artwork from Website", command=self.delete_artwork)
        menubar.add_cascade(label="Actions", menu=actions)
        self.master.config(menu=menubar)

        header = ttk.Frame(self.master, padding=(16, 10))
        header.pack(fill=tk.X)
        ttk.Label(header, text="🎨 Art Catalog Uploader", style="Header.TLabel").pack(side=tk.LEFT)
        ttk.Label(header, text="Search, manage, upload, edit, and generate descriptions.", style="Subtle.TLabel").pack(side=tk.LEFT, padx=(10, 0))

        paned = ttk.Panedwindow(self.master, orient=tk.HORIZONTAL)
        paned.pack(fill=tk.BOTH, expand=True, padx=16, pady=(0, 10))
        left = ttk.Frame(paned, padding=12)
        right = ttk.Frame(paned)
        paned.add(left, weight=0)
        paned.add(right, weight=1)

        ttk.Label(left, text="Using shared catalog database", style="Subtle.TLabel").pack(anchor="w", pady=(0, 10))
        ttk.Label(left, text="Search by Title or Artist").pack(anchor="w")
        search_row = ttk.Frame(left)
        search_row.pack(fill=tk.X, pady=(4, 12))
        self.search_var = tk.StringVar()
        self.search_entry = ttk.Entry(search_row, textvariable=self.search_var)
        self.search_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.search_entry.bind("<Return>", self.search_catalog)
        ttk.Button(search_row, text="Search", command=self.search_catalog, style="Accent.TButton").pack(side=tk.LEFT, padx=(8, 0))

        for text, cmd, style in [
            ("Export Search -> Word", self.export_to_word, None),
            ("Export Entire Database", self.export_entire_database, None),
            ("Upload Artwork(s) to Website", self.upload_artworks, "Success.TButton"),
            ("Generate Description for Selected", self.generate_description_for_selected, "Accent.TButton"),
            ("Edit Website Listings", self.edit_website_listings, "Accent.TButton"),
            ("Delete Artwork from Website", self.delete_artwork, "Danger.TButton"),
            ("Go to Your Listings", lambda: webbrowser.open(f"{BASE_URL}/artworks/"), None),
            ("Display All Inventory", self.display_all_inventory, "Accent.TButton"),
        ]:
            ttk.Button(left, text=text, command=cmd, style=style).pack(fill=tk.X, pady=4)

        ttk.Label(left, text="Shortcuts\nCtrl+F Search\nCtrl+E Export search\nCtrl+D Export DB\nCtrl+U Upload\nCtrl+Q Quit", style="Subtle.TLabel", justify="left").pack(anchor="w", pady=(16, 0))

        notebook = ttk.Notebook(right)
        notebook.pack(fill=tk.BOTH, expand=True)
        tab_results = ttk.Frame(notebook, padding=10)
        tab_details = ttk.Frame(notebook, padding=10)
        tab_preview = ttk.Frame(notebook, padding=10)
        tab_auction_search = ttk.Frame(notebook, padding=10)
        notebook.add(tab_results, text="Results")
        notebook.add(tab_details, text="Details")
        notebook.add(tab_preview, text="Preview Text")
        notebook.add(tab_auction_search, text="Auction Search")

        cols = ("Artist", "Title", "Year", "Medium", "Catalog #")
        self.tree = ttk.Treeview(tab_results, columns=cols, show="headings", style="Catalog.Treeview", selectmode="extended")
        for c in cols:
            self.tree.heading(c, text=c, command=partial(self._sort_tree, c, False))
            self.tree.column(c, width=130, stretch=True)
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb = ttk.Scrollbar(tab_results, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=vsb.set)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree.bind("<<TreeviewSelect>>", self._on_select_row)

        self.details_text = scrolledtext.ScrolledText(tab_details, wrap=tk.WORD, font=("Segoe UI", 11))
        self.details_text.pack(fill=tk.BOTH, expand=True)
        self.details_text.configure(state=tk.DISABLED)
        self.text_display = scrolledtext.ScrolledText(tab_preview, wrap=tk.WORD, font=("Segoe UI", 11))
        self.text_display.pack(fill=tk.BOTH, expand=True)
        self.text_display.configure(state=tk.DISABLED)
        self._build_auction_search_tab(tab_auction_search)
        self.status = ttk.Label(self.master, text="Loading shared catalog database...", anchor="w", padding=(16, 6))
        self.status.pack(fill=tk.X)

    def _build_auction_search_tab(self, tab):
        tab.columnconfigure(0, weight=1)
        tab.rowconfigure(3, weight=1)

        ttk.Label(tab, text="Upcoming Print Auction Web Research", style="Header.TLabel").grid(row=0, column=0, sticky="w")
        ttk.Label(
            tab,
            text="Find qualifying print and multiples sales in the next 3 or 7 days. Research runs securely through the SecondState server.",
            style="Subtle.TLabel",
            wraplength=850,
        ).grid(row=1, column=0, sticky="ew", pady=(2, 10))

        controls = ttk.LabelFrame(tab, text="Search criteria", padding=10)
        controls.grid(row=2, column=0, sticky="ew", pady=(0, 10))
        controls.columnconfigure(3, weight=1)

        self.auction_horizon_var = tk.IntVar(value=7)
        self.auction_minimum_var = tk.StringVar(value="10")
        self.auction_region_var = tk.StringVar()

        ttk.Label(controls, text="Horizon").grid(row=0, column=0, sticky="w", padx=(0, 8))
        ttk.Radiobutton(controls, text="3 days", variable=self.auction_horizon_var, value=3).grid(row=0, column=1, sticky="w")
        ttk.Radiobutton(controls, text="7 days", variable=self.auction_horizon_var, value=7).grid(row=0, column=2, sticky="w", padx=(0, 18))
        ttk.Label(controls, text="Minimum print lots in a mixed sale").grid(row=0, column=3, sticky="e", padx=(0, 8))
        ttk.Entry(controls, textvariable=self.auction_minimum_var, width=8).grid(row=0, column=4, sticky="w")

        ttk.Label(controls, text="Region (optional)").grid(row=1, column=0, sticky="w", pady=(8, 0), padx=(0, 8))
        ttk.Entry(controls, textvariable=self.auction_region_var).grid(row=1, column=1, columnspan=4, sticky="ew", pady=(8, 0))
        ttk.Label(controls, text="Additional instructions (optional)").grid(row=2, column=0, sticky="nw", pady=(8, 0), padx=(0, 8))
        self.auction_instructions_text = tk.Text(controls, height=3, wrap=tk.WORD, font=("Segoe UI", 10))
        self.auction_instructions_text.grid(row=2, column=1, columnspan=4, sticky="ew", pady=(8, 0))

        action_row = ttk.Frame(controls)
        action_row.grid(row=3, column=0, columnspan=5, sticky="ew", pady=(10, 0))
        action_row.columnconfigure(2, weight=1)
        self.auction_search_button = ttk.Button(
            action_row,
            text="Search",
            command=self.search_upcoming_auctions,
            style="Accent.TButton",
        )
        self.auction_search_button.grid(row=0, column=0, sticky="w")
        self.auction_progress = ttk.Progressbar(action_row, mode="indeterminate", length=150)
        self.auction_progress.grid(row=0, column=1, sticky="w", padx=(10, 0))
        self.auction_status = ttk.Label(action_row, text="Ready.", style="Subtle.TLabel")
        self.auction_status.grid(row=0, column=2, sticky="w", padx=(10, 0))

        output_frame = ttk.LabelFrame(tab, text="Markdown output", padding=8)
        output_frame.grid(row=3, column=0, sticky="nsew")
        self.auction_markdown_text = scrolledtext.ScrolledText(output_frame, wrap=tk.WORD, font=("Consolas", 10))
        self.auction_markdown_text.pack(fill=tk.BOTH, expand=True)

        output_buttons = ttk.Frame(tab)
        output_buttons.grid(row=4, column=0, sticky="ew", pady=(10, 0))
        ttk.Button(output_buttons, text="Copy Markdown", command=self.copy_auction_markdown).pack(side=tk.LEFT)
        ttk.Button(output_buttons, text="Save Markdown", command=self.save_auction_markdown).pack(side=tk.LEFT, padx=(8, 0))

    def _set_auction_search_busy(self, busy, status_text):
        self.auction_search_button.config(state=tk.DISABLED if busy else tk.NORMAL)
        self.auction_status.config(text=status_text)
        if busy:
            self.auction_progress.start(12)
        else:
            self.auction_progress.stop()

    def _request_upcoming_auction_search(self, payload):
        response = requests.post(
            f"{BASE_URL}/artworks/search_upcoming_print_auctions/",
            json=payload,
            headers=api_headers({"Content-Type": "application/json"}),
            timeout=(10, 930),
        )
        if response.status_code >= 400:
            try:
                detail = response.json().get("error") or response.text
            except ValueError:
                detail = response.text
            raise RuntimeError(detail or f"Auction search failed with status {response.status_code}.")
        try:
            result = response.json()
        except ValueError as exc:
            raise RuntimeError("The server returned an invalid auction-search response.") from exc
        if not isinstance(result.get("markdown"), str):
            raise RuntimeError("The server response did not include Markdown output.")
        return result

    def search_upcoming_auctions(self):
        try:
            minimum_count = int(self.auction_minimum_var.get().strip())
        except ValueError:
            messagebox.showwarning("Invalid Minimum", "Minimum print lots must be a whole number from 1 to 500.")
            return
        if not 1 <= minimum_count <= 500:
            messagebox.showwarning("Invalid Minimum", "Minimum print lots must be a whole number from 1 to 500.")
            return

        payload = {
            "horizon_days": self.auction_horizon_var.get(),
            "minimum_print_lots": minimum_count,
            "region": self.auction_region_var.get().strip(),
            "additional_instructions": self.auction_instructions_text.get("1.0", tk.END).strip(),
            "client_now": datetime.now().astimezone().isoformat(timespec="seconds"),
        }
        self.auction_markdown_text.delete("1.0", tk.END)
        self._set_auction_search_busy(True, "Researching official auction sources…")
        self._set_status("Auction web research is running.")

        def worker():
            try:
                result = self._request_upcoming_auction_search(payload)
            except Exception as exc:
                error_message = str(exc)
                self.master.after(0, lambda: self._finish_auction_search_error(error_message))
                return
            self.master.after(0, lambda: self._finish_auction_search_success(result))

        threading.Thread(target=worker, daemon=True).start()

    def _finish_auction_search_success(self, result):
        markdown = result.get("markdown", "")
        research_meta = result.get("research_meta") if isinstance(result.get("research_meta"), dict) else {}
        search_count = research_meta.get("search_count", 0)
        source_count = research_meta.get("source_count", 0)
        candidate_count = research_meta.get("raw_candidate_count", 0)
        qualified_count = research_meta.get("qualified_count", result.get("auction_count", 0))
        self.auction_markdown_text.delete("1.0", tk.END)
        self.auction_markdown_text.insert(tk.END, markdown)
        label = (
            f"Ran {search_count} searches and opened {source_count} sources; "
            f"found {candidate_count} candidates and {qualified_count} qualifying auctions."
        )
        self._set_auction_search_busy(False, label)
        self._set_status(label)

    def _finish_auction_search_error(self, error_message):
        self._set_auction_search_busy(False, "Search failed.")
        self._set_status("Auction web research failed.")
        messagebox.showerror("Auction Search Failed", error_message)

    def copy_auction_markdown(self):
        markdown = self.auction_markdown_text.get("1.0", tk.END).strip()
        if not markdown:
            messagebox.showwarning("Nothing to Copy", "Run an auction search first.")
            return
        self.master.clipboard_clear()
        self.master.clipboard_append(markdown)
        self.auction_status.config(text="Markdown copied to clipboard.")

    def save_auction_markdown(self):
        markdown = self.auction_markdown_text.get("1.0", tk.END).strip()
        if not markdown:
            messagebox.showwarning("Nothing to Save", "Run an auction search first.")
            return
        file_path = filedialog.asksaveasfilename(
            title="Save auction research",
            defaultextension=".md",
            initialfile=f"upcoming-print-auctions-{datetime.now():%Y-%m-%d}.md",
            filetypes=[("Markdown", "*.md"), ("Text", "*.txt"), ("All Files", "*.*")],
        )
        if not file_path:
            return
        with open(file_path, "w", encoding="utf-8", newline="\n") as output_file:
            output_file.write(markdown + "\n")
        self.auction_status.config(text=f"Saved {os.path.basename(file_path)}.")

    def _set_status(self, text):
        self.status.config(text=text)

    def safe(self, val, fallback=""):
        return str(val).strip() if pd.notna(val) else fallback

    def _sort_tree(self, col, reverse):
        data = [(self.tree.set(k, col), k) for k in self.tree.get_children("")]
        data.sort(key=lambda t: (t[0] or "").lower(), reverse=reverse)
        for index, (_val, k) in enumerate(data):
            self.tree.move(k, "", index)
        self.tree.heading(col, command=lambda c=col: self._sort_tree(c, not reverse))

    def _selected_catalog_row(self):
        sel = self.tree.selection()
        if not sel or self.current_results is None:
            return None
        return self.current_results.loc[int(sel[0])]

    def _dimensions_from_row(self, row):
        raw_height = str(row.get("Height", "")).strip()
        raw_width = str(row.get("Width", "")).strip()
        if raw_height.startswith("S"):
            return "", f'{raw_height.replace("S", "").strip()}" x {raw_width.replace("S", "").strip()}"'
        return f'{raw_height}" x {raw_width}"', ""

    def _row_to_payload(self, row, price="", catalog_description=""):
        dimensions_text, sheet_size = self._dimensions_from_row(row)
        return {
            "artist": self.safe(row.get("Artist", "")),
            "title": self.safe(row.get("Title", "")) or "Untitled",
            "year": self.safe(row.get("Year", "")),
            "medium": self.safe(row.get("Medium", "")),
            "description": self.safe(row.get("Description/Notes", "")),
            "dimensions_text": dimensions_text,
            "sheet_size": sheet_size,
            "catalog_number": self.safe(row.get("Catalog Number", "")),
            "price": price,
            "catalog_description": catalog_description,
        }

    def _load_excel_from_path(self, path):
        preview = pd.read_excel(path, sheet_name=DEFAULT_CATALOG_SHEET, header=None, nrows=30)
        header_row = None
        for i in range(len(preview)):
            row_vals = preview.iloc[i].astype(str).str.strip().tolist()
            if "Title" in row_vals and "Artist" in row_vals:
                header_row = i
                break
        if header_row is None:
            raise ValueError(f"Could not find a header row containing both 'Title' and 'Artist' on sheet '{DEFAULT_CATALOG_SHEET}'.")
        self.df = pd.read_excel(path, sheet_name=DEFAULT_CATALOG_SHEET, header=header_row)
        self.df.columns = [str(c).strip() for c in self.df.columns]
        self._clear_results()
        self._set_status(f"Loaded: {os.path.basename(path)} — {len(self.df):,} records")

    def load_excel(self):
        try:
            if not os.path.exists(DEFAULT_CATALOG_PATH):
                raise FileNotFoundError(f"Shared catalog file not found:\n{DEFAULT_CATALOG_PATH}")
            self._load_excel_from_path(DEFAULT_CATALOG_PATH)
        except Exception as e:
            messagebox.showerror("Load Error", f"Failed to load Excel file:\n{e}")

    def _clear_results(self):
        for i in self.tree.get_children():
            self.tree.delete(i)
        self.current_results = None
        for widget in (self.details_text, self.text_display):
            widget.configure(state=tk.NORMAL)
            widget.delete("1.0", tk.END)
            widget.configure(state=tk.DISABLED)

    def _fill_results(self, results):
        self._clear_results()
        self.current_results = results
        for idx, row in results.iterrows():
            self.tree.insert("", tk.END, iid=str(idx), values=(self.safe(row.get("Artist", "")), self.safe(row.get("Title", "")), self.safe(row.get("Year", "")), self.safe(row.get("Medium", "")), self.safe(row.get("Catalog Number", ""))))
        entries = [self.format_catalog_entry(row) for _, row in results.iterrows()]
        self.text_display.configure(state=tk.NORMAL)
        self.text_display.insert(tk.END, "\n\n".join(entries))
        self.text_display.configure(state=tk.DISABLED)

    def display_all_inventory(self):
        if self.df is None:
            messagebox.showwarning("No Catalog", "Open an Excel catalog first.")
            return
        self._fill_results(self.df)
        self._set_status(f"Showing all inventory: {len(self.df)} records.")

    def search_catalog(self, event=None):
        if self.df is None:
            messagebox.showwarning("No Catalog", "Open an Excel catalog first.")
            return
        query = self.search_var.get().lower().strip()
        if not query:
            messagebox.showwarning("Warning", "Enter a search query!")
            return
        results = self.df[self.df["Title"].astype(str).str.lower().str.contains(query, na=False) | self.df["Artist"].astype(str).str.lower().str.contains(query, na=False)]
        if results.empty:
            self._clear_results()
            self._set_status("No results found.")
            return
        self._fill_results(results)
        self._set_status(f"Found {len(results)} result(s).")

    def _on_select_row(self, event=None):
        row = self._selected_catalog_row()
        if row is None:
            return
        self.details_text.configure(state=tk.NORMAL)
        self.details_text.delete("1.0", tk.END)
        self.details_text.insert(tk.END, self.format_catalog_entry(row))
        self.details_text.configure(state=tk.DISABLED)

    def format_catalog_entry(self, row):
        artist = self.safe(row.get("Artist", "")).upper() or "UNKNOWN ARTIST"
        title = self.safe(row.get("Title", "")) or "Untitled"
        year = self.safe(row.get("Year", ""), "Unknown Year")
        medium = self.safe(row.get("Medium", ""), "Unknown Medium")
        notes = self.safe(row.get("Description/Notes", ""))
        dimensions_text, sheet_size = self._dimensions_from_row(row)
        size_label = "Sheet Size" if sheet_size else "Image Size"
        size_value = sheet_size or dimensions_text
        catalog_number = self.safe(row.get("Catalog Number", ""), "N/A")
        low = f"${int(row['Low']):,}" if pd.notna(row.get("Low")) else "N/A"
        high = f"${int(row['High']):,}" if pd.notna(row.get("High")) else "N/A"
        text = f"{artist}\n{title}, {year}\n{medium}"
        if notes:
            text += f"\n{notes}"
        text += f"\n{size_label}: {size_value}"
        if catalog_number:
            text += f"\nCatalog #: {catalog_number}"
        return text + f"\n\nEstimate: {low} - {high}"

    def _fetch_website_artworks(self):
        response = requests.get(f"{BASE_URL}/artworks/manage_json/", headers=api_headers(), timeout=25)
        response.raise_for_status()
        return response.json().get("artworks", [])

    def _generate_description(self, payload):
        response = requests.post(f"{BASE_URL}/artworks/generate_description/", json=payload, headers=api_headers({"Content-Type": "application/json"}), timeout=70)
        response.raise_for_status()
        description = response.json().get("description", "").strip()
        if not description:
            raise ValueError("Website returned an empty generated description.")
        return description

    def _description_dialog(self, title, payload, initial_text=""):
        dialog = tk.Toplevel(self.master)
        dialog.title(title)
        dialog.geometry("760x620")
        dialog.transient(self.master)
        dialog.grab_set()
        frm = ttk.Frame(dialog, padding=16)
        frm.pack(fill=tk.BOTH, expand=True)
        ttk.Label(frm, text=title, style="Header.TLabel").pack(anchor="w")
        text = scrolledtext.ScrolledText(frm, wrap=tk.WORD, height=18, font=("Segoe UI", 11))
        text.pack(fill=tk.BOTH, expand=True, pady=(10, 0))
        text.insert(tk.END, initial_text or payload.get("catalog_description", ""))
        use_web = tk.BooleanVar(value=True)
        status = ttk.Label(frm, text="", style="Subtle.TLabel")
        status.pack(anchor="w", pady=(8, 0))
        result = {"value": None}

        def generate():
            generate_btn.config(state=tk.DISABLED)
            try:
                working = dict(payload)
                working["catalog_description"] = text.get("1.0", tk.END).strip()
                working["use_web"] = use_web.get()
                status.config(text="Generating description...")
                dialog.update_idletasks()
                description = self._generate_description(working)
                text.delete("1.0", tk.END)
                text.insert(tk.END, description)
                status.config(text="Draft inserted. Review/edit before using it.")
            except Exception as exc:
                messagebox.showerror("Generate Description Failed", str(exc), parent=dialog)
            finally:
                generate_btn.config(state=tk.NORMAL)

        def accept():
            result["value"] = text.get("1.0", tk.END).strip()
            dialog.destroy()

        btns = ttk.Frame(frm)
        btns.pack(fill=tk.X, pady=(10, 0))
        generate_btn = ttk.Button(btns, text="Generate with OpenAI", command=generate, style="Accent.TButton")
        generate_btn.pack(side=tk.LEFT)
        ttk.Checkbutton(btns, text="allow web search", variable=use_web).pack(side=tk.LEFT, padx=(10, 0))
        ttk.Button(btns, text="Use Description", command=accept, style="Success.TButton").pack(side=tk.RIGHT)
        ttk.Button(btns, text="Cancel", command=dialog.destroy).pack(side=tk.RIGHT, padx=(0, 8))
        dialog.wait_window()
        return result["value"]

    def generate_description_for_selected(self):
        row = self._selected_catalog_row()
        if row is None:
            messagebox.showwarning("No Selection", "Select one artwork row first.")
            return
        payload = self._row_to_payload(row)
        description = self._description_dialog(f"Description for {payload['title']}", payload)
        if description is not None:
            self.text_display.configure(state=tk.NORMAL)
            self.text_display.delete("1.0", tk.END)
            self.text_display.insert(tk.END, description)
            self.text_display.configure(state=tk.DISABLED)
            self._set_status("Generated description is shown in Preview Text. It is not saved until upload/edit save.")

    def edit_website_listings(self):
        try:
            artworks = self._fetch_website_artworks()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to fetch artworks: {e}")
            return
        if not artworks:
            messagebox.showinfo("No Listings", "No website listings found.")
            return
        win = tk.Toplevel(self.master)
        win.title("Edit Website Listings")
        win.geometry("1180x760")
        paned = ttk.Panedwindow(win, orient=tk.HORIZONTAL)
        paned.pack(fill=tk.BOTH, expand=True, padx=14, pady=14)
        left = ttk.Frame(paned, padding=8)
        right = ttk.Frame(paned, padding=8)
        paned.add(left, weight=1)
        paned.add(right, weight=2)
        ttk.Label(left, text="Website Listings", style="Header.TLabel").pack(anchor="w")
        listing_tree = ttk.Treeview(left, columns=("Artist", "Title", "Price"), show="headings")
        for col in ("Artist", "Title", "Price"):
            listing_tree.heading(col, text=col)
            listing_tree.column(col, width=140, stretch=True)
        listing_tree.pack(fill=tk.BOTH, expand=True, pady=(8, 0))
        for art in artworks:
            listing_tree.insert("", tk.END, iid=str(art.get("id")), values=(art.get("artist", ""), art.get("title", ""), art.get("price", "")))

        fields, selected = {}, {"art": None}
        is_available = tk.BooleanVar(value=True)
        form = ttk.Frame(right)
        form.pack(fill=tk.BOTH, expand=True)
        form.columnconfigure(1, weight=1)
        row_num = 0
        ttk.Label(form, text="Edit Listing", style="Header.TLabel").grid(row=row_num, column=0, columnspan=2, sticky="w")
        row_num += 1
        for key, label in self.EDIT_FIELDS:
            ttk.Label(form, text=label).grid(row=row_num, column=0, sticky="w", padx=(0, 8), pady=3)
            entry = ttk.Entry(form)
            entry.grid(row=row_num, column=1, sticky="ew", pady=3)
            fields[key] = entry
            row_num += 1
        ttk.Label(form, text="Notes / signature text").grid(row=row_num, column=0, sticky="nw", padx=(0, 8), pady=3)
        notes_text = scrolledtext.ScrolledText(form, wrap=tk.WORD, height=4, font=("Segoe UI", 10))
        notes_text.grid(row=row_num, column=1, sticky="ew", pady=3)
        row_num += 1
        ttk.Label(form, text="Description").grid(row=row_num, column=0, sticky="nw", padx=(0, 8), pady=3)
        desc_text = scrolledtext.ScrolledText(form, wrap=tk.WORD, height=8, font=("Segoe UI", 10))
        desc_text.grid(row=row_num, column=1, sticky="nsew", pady=3)
        form.rowconfigure(row_num, weight=1)
        row_num += 1
        ttk.Checkbutton(form, text="Available", variable=is_available).grid(row=row_num, column=1, sticky="w")
        row_num += 1
        image_box = ttk.LabelFrame(form, text="Images")
        image_box.grid(row=row_num, column=0, columnspan=2, sticky="ew", pady=(8, 4))
        image_list = ttk.Frame(image_box, padding=6)
        image_list.pack(fill=tk.X)
        image_status = ttk.Label(image_box, text="", style="Subtle.TLabel")
        image_status.pack(anchor="w", padx=6, pady=(0, 6))
        delete_image_vars, add_image_paths = {}, []
        status = ttk.Label(right, text="Select a listing to edit.", style="Subtle.TLabel")
        status.pack(anchor="w", pady=(8, 0))

        def payload():
            data = {key: widget.get().strip() for key, widget in fields.items()}
            data["description"] = notes_text.get("1.0", tk.END).strip()
            data["catalog_description"] = desc_text.get("1.0", tk.END).strip()
            data["is_available"] = is_available.get()
            return data

        def populate(art):
            selected["art"] = art
            for key, widget in fields.items():
                widget.delete(0, tk.END)
                widget.insert(0, art.get(key, "") or "")
            notes_text.delete("1.0", tk.END)
            notes_text.insert(tk.END, art.get("description", "") or "")
            desc_text.delete("1.0", tk.END)
            desc_text.insert(tk.END, art.get("catalog_description", "") or "")
            is_available.set(bool(art.get("is_available", True)))
            add_image_paths.clear()
            for child in image_list.winfo_children():
                child.destroy()
            delete_image_vars.clear()
            for image in art.get("images", []) or []:
                var = tk.BooleanVar(value=False)
                delete_image_vars[str(image.get("id"))] = var
                ttk.Checkbutton(image_list, text=f"Delete image {image.get('id')}: {image.get('url', '')}", variable=var).pack(anchor="w", fill=tk.X)
            image_status.config(text="")
            status.config(text=f"Editing listing ID {art.get('id')}")

        def on_select(_event=None):
            sel = listing_tree.selection()
            if sel:
                art = next((a for a in artworks if str(a.get("id")) == sel[0]), None)
                if art:
                    populate(art)

        def add_images():
            paths = filedialog.askopenfilenames(title="Add images", filetypes=[("Image Files", "*.jpg *.jpeg *.png")], parent=win)
            add_image_paths.extend(paths)
            image_status.config(text=f"{len(add_image_paths)} new image(s) queued.")

        def generate():
            art = selected["art"]
            if not art:
                return
            try:
                data = payload()
                data["use_web"] = True
                response = requests.post(f"{BASE_URL}/artworks/{art.get('id')}/generate_description/", json=data, headers=api_headers({"Content-Type": "application/json"}), timeout=70)
                response.raise_for_status()
                desc_text.delete("1.0", tk.END)
                desc_text.insert(tk.END, response.json().get("description", ""))
                status.config(text="Draft inserted. Review/edit, then Save Changes.")
            except Exception as exc:
                messagebox.showerror("Generate Failed", str(exc), parent=win)

        def save():
            art = selected["art"]
            if not art:
                return
            data = payload()
            for image_id, var in delete_image_vars.items():
                if var.get():
                    data.setdefault("delete_image_ids", []).append(image_id)
            data_items = []
            for key, value in data.items():
                if isinstance(value, list):
                    data_items.extend((key, item) for item in value)
                else:
                    data_items.append((key, str(value)))
            files = []
            try:
                for i, path in enumerate(add_image_paths):
                    mime = "image/png" if path.lower().endswith(".png") else "image/jpeg"
                    files.append((f"image_{i}", (os.path.basename(path), open(path, "rb"), mime)))
                response = requests.post(f"{BASE_URL}/artworks/{art.get('id')}/update_artwork/", data=data_items, files=files, headers=api_headers(), timeout=70)
                response.raise_for_status()
                updated = response.json().get("artwork", {})
                listing_tree.item(str(art.get("id")), values=(updated.get("artist", ""), updated.get("title", ""), updated.get("price", "")))
                populate(updated)
                status.config(text="Saved changes to website listing.")
            except Exception as exc:
                messagebox.showerror("Save Failed", str(exc), parent=win)
            finally:
                for _name, file_tuple in files:
                    try:
                        file_tuple[1].close()
                    except Exception:
                        pass

        listing_tree.bind("<<TreeviewSelect>>", on_select)
        buttons = ttk.Frame(right)
        buttons.pack(fill=tk.X, pady=(10, 0))
        ttk.Button(buttons, text="Generate Description", command=generate, style="Accent.TButton").pack(side=tk.LEFT)
        ttk.Button(buttons, text="Add Images", command=add_images).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Button(buttons, text="Save Changes", command=save, style="Success.TButton").pack(side=tk.LEFT, padx=(8, 0))
        ttk.Button(buttons, text="Close", command=win.destroy).pack(side=tk.RIGHT)
        first = str(artworks[0].get("id"))
        listing_tree.selection_set(first)
        listing_tree.focus(first)
        populate(artworks[0])

    def export_to_word(self):
        text = self.text_display.get("1.0", tk.END).strip()
        if not text:
            messagebox.showwarning("Nothing to Export", "Search or generate text first.")
            return
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not file_path:
            return
        doc = Document()
        for entry in text.split("\n\n"):
            doc.add_paragraph(entry)
        doc.save(file_path)
        self._set_status(f"Exported: {os.path.basename(file_path)}")

    def export_entire_database(self):
        if self.df is None:
            messagebox.showwarning("No Catalog", "Open an Excel catalog first.")
            return
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not file_path:
            return
        doc = Document()
        for _, row in self.df.iterrows():
            doc.add_paragraph(self.format_catalog_entry(row))
        doc.save(file_path)
        self._set_status(f"Exported: {os.path.basename(file_path)}")

    def delete_artwork(self):
        messagebox.showinfo("Moved", "Use Edit Website Listings for editing/images, or existing website deletion flow if needed.")

    def upload_artworks(self):
        if self.current_results is None or self.current_results.empty:
            messagebox.showwarning("Warning", "No artworks available. Search or display inventory first.")
            return
        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showwarning("No Selection", "Select one or more artworks first.")
            return
        for _, row in self.current_results.loc[[int(iid) for iid in selected_items]].iterrows():
            title = row.get("Title", "")
            price = simpledialog.askstring("Enter Sale Price", f"Enter sale price for '{title}':")
            if not price:
                continue
            data = self._row_to_payload(row, price=price)
            description = self._description_dialog(f"Description for upload: {data['title']}", data)
            if description is None:
                continue
            data["catalog_description"] = description
            images = list(filedialog.askopenfilenames(title=f"Select image(s) for '{title}'", filetypes=[("Image Files", "*.jpg *.jpeg *.png")]))
            if not images:
                continue
            files = []
            try:
                for i, img in enumerate(images):
                    mime = "image/png" if img.lower().endswith(".png") else "image/jpeg"
                    files.append((f"image_{i}", (os.path.basename(img), open(img, "rb"), mime)))
                response = requests.post(f"{BASE_URL}/artworks/upload_artwork/", data=data, files=files, headers=api_headers(), timeout=70)
                if response.status_code == 201:
                    messagebox.showinfo("Success", f"'{title}' uploaded successfully!")
                else:
                    messagebox.showerror("Error", f"Upload failed: {response.text}")
            finally:
                for _name, file_tuple in files:
                    try:
                        file_tuple[1].close()
                    except Exception:
                        pass


def main():
    root = tk.Tk()
    ArtCatalogApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
