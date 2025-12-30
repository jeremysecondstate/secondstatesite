import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk
from tkinter import scrolledtext
import pandas as pd
import requests
from docx import Document
import webbrowser
import os
from functools import partial


# ----- App Config -----
BASE_URL = "https://secondstate.art"
APP_TITLE = "Art Catalog Uploader"
APP_MIN_W, APP_MIN_H = 1100, 780
CATALOG_API_KEY = "276e19f127f140623e73e6c160bbd8ed"
SALES_SHEET_ENDPOINT = f"{BASE_URL}/sales/upload_sales_sheet/"

def api_headers():
    return {"X-API-KEY": CATALOG_API_KEY}


class ArtCatalogApp:
    def __init__(self, master):
        self.master = master
        master.title(APP_TITLE)
        master.geometry(f"{APP_MIN_W}x{APP_MIN_H}")
        master.minsize(APP_MIN_W, APP_MIN_H)

        # -------- THEME / STYLE --------
        self.style = ttk.Style(master)
        # Choose a reliable base theme and customize
        try:
            self.style.theme_use("clam")
        except tk.TclError:
            pass

        ACCENT = "#4E8DF5"
        SUCCESS = "#28a745"
        DANGER = "#d9534f"
        BG = "#F6F7FB"

        master.configure(bg=BG)

        # General paddings & fonts
        self.style.configure(".", font=("Segoe UI", 11))
        self.style.configure("Header.TLabel", font=("Segoe UI Semibold", 16))
        self.style.configure("Subtle.TLabel", foreground="#5a5f73")
        self.style.configure("TButton", padding=8)
        self.style.map(
            "TButton",
            foreground=[("active", "!disabled", "#000")],
            background=[("active", "!disabled", "#E8ECF7")],
        )
        self.style.configure("Accent.TButton", padding=10, font=("Segoe UI Semibold", 11), foreground="white", background=ACCENT)
        self.style.map("Accent.TButton",
                       background=[("active", "!disabled", "#3E78D1")])
        self.style.configure("Success.TButton", padding=10, font=("Segoe UI Semibold", 11), foreground="white", background=SUCCESS)
        self.style.map("Success.TButton",
                       background=[("active", "!disabled", "#1f7e37")])
        self.style.configure("Danger.TButton", padding=10, font=("Segoe UI Semibold", 11), foreground="white", background=DANGER)
        self.style.map("Danger.TButton",
                       background=[("active", "!disabled", "#b8453f")])

        # Treeview styling
        self.style.configure("Catalog.Treeview", rowheight=28, borderwidth=0, relief="flat")
        self.style.configure("Catalog.Treeview.Heading", font=("Segoe UI Semibold", 11))
        self.style.map("Catalog.Treeview",
                       background=[("selected", "#E3EDFF")],
                       foreground=[("selected", "#000")])

        # ------- State -------
        self.df = None
        self.current_results = None
        self.open_listings_button = None
        self.mode = "inventory"  # "inventory" or "sales"
        self.sales_path = None  # remember the sales xlsx path for upload

        # ------- MENUBAR -------
        self._build_menubar()

        # ------- HEADER -------
        header = ttk.Frame(master, padding=(16, 10, 16, 10))
        header.pack(fill=tk.X)
        ttk.Label(header, text="🎨 Art Catalog Uploader", style="Header.TLabel").pack(side=tk.LEFT)
        ttk.Label(header, text="Search, manage, and upload artworks to your website.", style="Subtle.TLabel").pack(side=tk.LEFT, padx=(10, 0))

        # ------- MAIN LAYOUT: Paned Window ------
        paned = ttk.Panedwindow(master, orient=tk.HORIZONTAL)
        paned.pack(fill=tk.BOTH, expand=True, padx=16, pady=(0, 10))

        # Left: Controls
        left = ttk.Frame(paned, padding=12)
        left.configure(style="Card.TFrame")
        paned.add(left, weight=0)

        # Right: Results
        right = ttk.Frame(paned, padding=0)
        paned.add(right, weight=1)

        # ------- LEFT / CONTROL CARD -------
        # Data loader row
        loader_row = ttk.Frame(left)
        loader_row.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(loader_row, text="Catalog Source:", style="Subtle.TLabel").pack(side=tk.LEFT)

        btns = ttk.Frame(loader_row)
        btns.pack(side=tk.RIGHT)

        ttk.Button(btns, text="Open Inventory Excel…", command=self.load_inventory_excel, style="Accent.TButton").pack(
            side=tk.LEFT, padx=(0, 8))
        ttk.Button(btns, text="Open Sales Excel…", command=self.load_sales_excel).pack(side=tk.LEFT)

        # Search box
        ttk.Label(left, text="Search by Title or Artist").pack(anchor="w", pady=(8, 4))
        self.search_var = tk.StringVar()
        search_row = ttk.Frame(left)
        search_row.pack(fill=tk.X)
        self.search_entry = ttk.Entry(search_row, textvariable=self.search_var)
        self.search_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.search_entry.bind("<Return>", self.search_catalog)
        ttk.Button(search_row, text="Search", command=self.search_catalog, style="Accent.TButton").pack(side=tk.LEFT, padx=(8, 0))

        # Actions group
        ttk.Label(left, text="Actions").pack(anchor="w", pady=(16, 6))
        ttk.Button(left, text="Export Search → Word", command=self.export_to_word).pack(fill=tk.X, pady=4)
        ttk.Button(left, text="Export Entire Database", command=self.export_entire_database).pack(fill=tk.X, pady=4)
        ttk.Button(left, text="Upload Artwork(s) to Website", command=self.upload_artworks, style="Success.TButton").pack(fill=tk.X, pady=(10, 4))
        ttk.Button(left, text="Delete Artwork from Website", command=self.delete_artwork, style="Danger.TButton").pack(fill=tk.X, pady=4)
        ttk.Separator(left).pack(fill=tk.X, pady=12)
        ttk.Button(left, text="Go to Your Listings", command=lambda: webbrowser.open(f"{BASE_URL}/artworks/")).pack(fill=tk.X, pady=4)
        ttk.Button(left, text="Upload Sales Spreadsheet → Pieces Sold", command=self.upload_sales_sheet).pack(fill=tk.X,
                                                                                                              pady=(6,
                                                                                                                    4))

        # Helpful shortcuts info
        ttk.Label(left, text="Shortcuts", style="Subtle.TLabel").pack(anchor="w", pady=(16, 6))
        shortcuts = (
            "Ctrl+F: Search\n"
            "Ctrl+E: Export search\n"
            "Ctrl+D: Export database\n"
            "Ctrl+U: Upload results\n"
            "Ctrl+Q: Quit"
        )
        ttk.Label(left, text=shortcuts, justify="left").pack(anchor="w")

        # ------- RIGHT / RESULTS NOTEBOOK -------
        notebook = ttk.Notebook(right)
        notebook.pack(fill=tk.BOTH, expand=True)

        tab_results = ttk.Frame(notebook, padding=10)
        tab_details = ttk.Frame(notebook, padding=10)
        tab_preview = ttk.Frame(notebook, padding=10)
        notebook.add(tab_results, text="Results")
        notebook.add(tab_details, text="Details")
        notebook.add(tab_preview, text="Preview Text")

        # Results (Treeview)
        tree_frame = ttk.Frame(tab_results)
        tree_frame.pack(fill=tk.BOTH, expand=True)

        cols = ("Artist", "Title", "Year", "Medium", "Catalog #")
        self.tree = ttk.Treeview(tree_frame, columns=cols, show="headings", style="Catalog.Treeview")
        for c in cols:
            self.tree.heading(c, text=c, command=partial(self._sort_tree, c, False))
            self.tree.column(c, width=120, stretch=True)
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=vsb.set)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)

        self.tree.bind("<<TreeviewSelect>>", self._on_select_row)

        # Details panel
        self.details_text = scrolledtext.ScrolledText(tab_details, wrap=tk.WORD, height=20, font=("Segoe UI", 11))
        self.details_text.pack(fill=tk.BOTH, expand=True)
        self.details_text.configure(state=tk.DISABLED)

        # Preview panel (what gets exported by "Export Search → Word")
        self.text_display = scrolledtext.ScrolledText(tab_preview, wrap=tk.WORD, height=20, font=("Segoe UI", 11))
        self.text_display.pack(fill=tk.BOTH, expand=True)
        self.text_display.configure(state=tk.DISABLED)

        # ------- STATUS BAR -------
        self.status = ttk.Label(master, text="Open an Excel catalog to get started.", anchor="w", padding=(16, 6))
        self.status.pack(fill=tk.X)

        # ------- TOPMOST NUDGE (like original) -------
        self.master.lift()
        self.master.attributes("-topmost", True)
        self.master.after(600, lambda: self.master.attributes("-topmost", False))

        # ------- KEYBOARD SHORTCUTS -------
        master.bind_all("<Control-f>", self.search_catalog)
        master.bind_all("<Control-F>", self.search_catalog)
        master.bind_all("<Control-e>", lambda e: self.export_to_word())
        master.bind_all("<Control-E>", lambda e: self.export_to_word())
        master.bind_all("<Control-d>", lambda e: self.export_entire_database())
        master.bind_all("<Control-D>", lambda e: self.export_entire_database())
        master.bind_all("<Control-u>", lambda e: self.upload_artworks())
        master.bind_all("<Control-U>", lambda e: self.upload_artworks())
        master.bind_all("<Control-q>", lambda e: master.quit())
        master.bind_all("<Control-Q>", lambda e: master.quit())

        # Optional: auto-load last opened file if you like (commented)
        # self._try_load_last_excel()

    # -------------- UI Helpers --------------
    def _build_menubar(self):
        menubar = tk.Menu(self.master)
        file_menu = tk.Menu(menubar, tearoff=False)
        file_menu.add_command(label="Open Inventory Excel…", accelerator="Ctrl+O", command=self.load_inventory_excel)
        file_menu.add_separator()
        file_menu.add_command(label="Export Search → Word", accelerator="Ctrl+E", command=self.export_to_word)
        file_menu.add_command(label="Export Entire Database", accelerator="Ctrl+D", command=self.export_entire_database)
        file_menu.add_separator()
        file_menu.add_command(label="Quit", accelerator="Ctrl+Q", command=self.master.quit)
        menubar.add_cascade(label="File", menu=file_menu)

        actions_menu = tk.Menu(menubar, tearoff=False)
        actions_menu.add_command(label="Upload Artwork(s)", accelerator="Ctrl+U", command=self.upload_artworks)
        actions_menu.add_command(label="Delete Artwork from Website", command=self.delete_artwork)
        menubar.add_cascade(label="Actions", menu=actions_menu)

        help_menu = tk.Menu(menubar, tearoff=False)
        help_menu.add_command(label="Go to Listings", command=lambda: webbrowser.open(f"{BASE_URL}/artworks/"))
        menubar.add_cascade(label="Help", menu=help_menu)

        self.master.config(menu=menubar)
        # Bind Ctrl+O after menu creation
        self.master.bind_all("<Control-o>", lambda e: self.load_inventory_excel())
        self.master.bind_all("<Control-O>", lambda e: self.load_inventory_excel())

    def _set_status(self, text):
        self.status.config(text=text)

    def _sort_tree(self, col, reverse):
        try:
            data = [(self.tree.set(k, col), k) for k in self.tree.get_children("")]
            # Try numeric sort for Year; fallback to string
            if col == "Year":
                def to_int(x):
                    try:
                        return int(x[0])
                    except Exception:
                        return float("inf")
                data.sort(key=to_int, reverse=reverse)
            else:
                data.sort(key=lambda t: (t[0] or "").lower(), reverse=reverse)
            for index, (val, k) in enumerate(data):
                self.tree.move(k, "", index)
            self.tree.heading(col, command=lambda c=col: self._sort_tree(c, not reverse))
        except Exception as e:
            self._set_status(f"Sort error: {e}")

    def _on_select_row(self, event=None):
        sel = self.tree.selection()
        if not sel or self.current_results is None:
            return

        idx = int(sel[0])  # Treeview iid == dataframe index
        row = self.current_results.loc[idx]
        details = self.format_catalog_entry(row)

        self.details_text.configure(state=tk.NORMAL)
        self.details_text.delete("1.0", tk.END)
        self.details_text.insert(tk.END, details)
        self.details_text.configure(state=tk.DISABLED)

    # -------------- Data Loading --------------
    def load_inventory_excel(self):
        path = filedialog.askopenfilename(
            title="Select Inventory Excel Catalog",
            filetypes=[("Excel files", "*.xlsx *.xls")]
        )
        if not path:
            return

        try:
            preview = pd.read_excel(path, header=None, nrows=30)

            header_row = None
            for i in range(len(preview)):
                row_vals = preview.iloc[i].astype(str).str.strip().tolist()
                if "Title" in row_vals and "Artist" in row_vals:
                    header_row = i
                    break

            if header_row is None:
                raise ValueError("Could not find a header row containing both 'Title' and 'Artist'.")

            self.df = pd.read_excel(path, header=header_row)
            self.df.columns = [str(c).strip() for c in self.df.columns]

            self.mode = "inventory"
            self.sales_path = None  # not relevant here

            # Persist last path (optional)
            try:
                with open(os.path.join(os.path.dirname(__file__), ".last_catalog_path.txt"), "w",
                          encoding="utf-8") as f:
                    f.write(path)
            except Exception:
                pass

            self._clear_results()
            self._set_status(f"[Inventory] Loaded: {os.path.basename(path)} — {len(self.df):,} records")

        except Exception as e:
            messagebox.showerror("Load Error", f"Failed to load Excel file:\n{e}")

    def _try_load_last_excel(self):
        try:
            cfg = os.path.join(os.path.dirname(__file__), ".last_catalog_path.txt")
            if os.path.exists(cfg):
                with open(cfg, "r", encoding="utf-8") as f:
                    path = f.read().strip()
                if path and os.path.exists(path):
                    self.df = pd.read_excel(path)
                    self._set_status(f"Loaded: {os.path.basename(path)}  —  {len(self.df):,} records")
        except Exception:
            pass

    def load_sales_excel(self):
        path = filedialog.askopenfilename(
            title="Select Sales Excel Spreadsheet",
            filetypes=[("Excel files", "*.xlsx *.xls")]
        )
        if not path:
            return

        try:
            preview = pd.read_excel(path, header=None, nrows=30)

            header_row = None
            for i in range(len(preview)):
                row_vals = preview.iloc[i].astype(str).str.strip().tolist()
                # Common columns you showed earlier
                if "Artist" in row_vals and ("Name" in row_vals or "Title" in row_vals):
                    header_row = i
                    break

            if header_row is None:
                raise ValueError("Could not find a header row containing 'Artist' and 'Name' (or 'Title').")

            self.df = pd.read_excel(path, header=header_row)
            self.df.columns = [str(c).strip() for c in self.df.columns]

            # Normalize: treat Name as Title for searching/preview
            if "Title" not in self.df.columns and "Name" in self.df.columns:
                self.df["Title"] = self.df["Name"]

            self.mode = "sales"
            self.sales_path = path

            self._clear_results()
            self._set_status(f"[Sales] Loaded: {os.path.basename(path)} — {len(self.df):,} rows")

        except Exception as e:
            messagebox.showerror("Load Error", f"Failed to load Sales spreadsheet:\n{e}")

    def _clear_results(self):
        for i in self.tree.get_children():
            self.tree.delete(i)
        self.current_results = None
        self.details_text.configure(state=tk.NORMAL)
        self.details_text.delete("1.0", tk.END)
        self.details_text.configure(state=tk.DISABLED)
        self.text_display.configure(state=tk.NORMAL)
        self.text_display.delete("1.0", tk.END)
        self.text_display.configure(state=tk.DISABLED)

    # -------------- Core Formatting --------------
    def safe(self, val, fallback=""):
        return str(val).strip() if pd.notna(val) else fallback

    def format_catalog_entry(self, row):
        if self.mode == "sales":
            artist = self.safe(row.get("Artist", "")).upper() or "UNKNOWN ARTIST"
            title = self.safe(row.get("Title", "")) or "Untitled"
            date = self.safe(row.get("Date", ""), "Unknown Date")
            loc = self.safe(row.get("Sale Location", ""))
            price = self.safe(row.get("Net Sale Price $", ""))
            house = self.safe(row.get("Auction House", ""))

            out = f"{artist}\n{title}\nSOLD — {date}"
            if price:
                out += f"\nNet: {price}"
            if house or loc:
                out += f"\n{house}{' — ' if (house and loc) else ''}{loc}"
            return out

        # inventory mode (your existing formatting)
        artist = self.safe(row.get("Artist", "")).upper() or "UNKNOWN ARTIST"
        title = self.safe(row.get("Title", "")) or "Untitled"
        year = self.safe(row.get("Year", ""), "Unknown Year")
        medium = self.safe(row.get("Medium", ""), "Unknown Medium")
        description = self.safe(row.get("Description/Notes", ""))

        height = self.safe(row.get("Height", ""), "?")
        width = self.safe(row.get("Width", ""), "?")
        dimensions_text = f"{height}\" x {width}\""

        catalog_number = self.safe(row.get("Catalog Number", ""), "N/A")
        low_estimate = f"${int(row['Low']):,}" if pd.notna(row.get("Low")) else "N/A"
        high_estimate = f"${int(row['High']):,}" if pd.notna(row.get("High")) else "N/A"
        estimate_text = f"\n\nEstimate: {low_estimate} - {high_estimate}"

        formatted = f"{artist}\n{title}, {year}\n{medium}"
        if description:
            formatted += f"\n{description}"
        formatted += f"\nImage Size: {dimensions_text}"
        if catalog_number:
            formatted += f"\nCatalog #: {catalog_number}"
        formatted += estimate_text
        return formatted

    # -------------- Search --------------
    def search_catalog(self, event=None):
        if self.df is None:
            messagebox.showwarning("No Catalog", "Open an Excel file first.")
            return

        query = self.search_var.get().lower().strip()
        if not query:
            messagebox.showwarning("Warning", "Enter a search query!")
            return

        # In both modes, search "Title" + "Artist"
        results = self.df[
            self.df.get("Title", pd.Series(dtype=str)).astype(str).str.lower().str.contains(query, na=False) |
            self.df.get("Artist", pd.Series(dtype=str)).astype(str).str.lower().str.contains(query, na=False)
            ]

        self._clear_results()

        if results.empty:
            self._set_status("No results found.")
            self.text_display.configure(state=tk.NORMAL)
            self.text_display.insert(tk.END, "No results found.\n")
            self.text_display.configure(state=tk.DISABLED)
            return

        self.current_results = results

        # Populate treeview
        for idx, row in results.iterrows():
            self.tree.insert(
                "",
                tk.END,
                iid=str(idx),
                values=(
                    self.safe(row.get("Artist", "")),
                    self.safe(row.get("Title", "")),
                    self.safe(row.get("Year", "")) if self.mode == "inventory" else self.safe(row.get("Date", "")),
                    self.safe(row.get("Medium", "")) if self.mode == "inventory" else self.safe(
                        row.get("Net Sale Price $", "")),
                    self.safe(row.get("Catalog Number", "")) if self.mode == "inventory" else self.safe(
                        row.get("Auction House", "")),
                )
            )

        # Preview formatting
        entries = [self.format_catalog_entry(row) for _, row in results.iterrows()]
        self.text_display.configure(state=tk.NORMAL)
        self.text_display.insert(tk.END, "\n\n".join(entries))
        self.text_display.configure(state=tk.DISABLED)

        self._set_status(f"Found {len(results)} result(s). Mode: {self.mode}. Select a row to view details.")

    # -------------- Export --------------
    def export_to_word(self):
        text = self.text_display.get("1.0", tk.END).strip()
        if not text or text == "No results found.":
            messagebox.showwarning("Nothing to Export", "Search first, or select text to export.")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not file_path:
            return

        try:
            doc = Document()
            for entry in text.split("\n\n"):
                lines = entry.strip().split("\n")
                if not lines or not lines[0].strip():
                    continue

                artist_p = doc.add_paragraph()
                artist_p.add_run(lines[0]).bold = True

                if len(lines) > 1:
                    title_p = doc.add_paragraph()
                    title_p.add_run(lines[1]).italic = True

                for line in lines[2:]:
                    doc.add_paragraph(line)

                doc.add_paragraph("\n")

            doc.save(file_path)
            messagebox.showinfo("Success", "Exported search results successfully.")
            self._set_status(f"Exported: {os.path.basename(file_path)}")
        except Exception as e:
            messagebox.showerror("Export Error", str(e))

    def export_entire_database(self):
        if self.df is None:
            messagebox.showwarning("No Catalog", "Open an Excel catalog first.")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not file_path:
            return

        try:
            doc = Document()
            for _, row in self.df.iterrows():
                artist = str(row.get("Artist", "")).strip().upper()
                title = str(row.get("Title", "")).strip()
                year = str(row.get("Year", "")).strip()

                p = doc.add_paragraph()
                p.add_run(artist).bold = True
                p.add_run("\n")
                tp = doc.add_paragraph()
                tp.add_run(title).italic = True
                if year:
                    tp.add_run(f", {year}")

                doc.add_paragraph(self.format_catalog_entry(row))
                doc.add_paragraph("\n")

            doc.save(file_path)
            messagebox.showinfo("Success", "Database exported successfully.")
            self._set_status(f"Exported: {os.path.basename(file_path)}")
        except Exception as e:
            messagebox.showerror("Export Error", str(e))

    # -------------- Delete --------------
    def delete_artwork(self):
        delete_window = tk.Toplevel(self.master)
        delete_window.title("Delete Artwork")
        delete_window.geometry("520x220")
        frm = ttk.Frame(delete_window, padding=16)
        frm.pack(fill=tk.BOTH, expand=True)

        ttk.Label(frm, text="Select Artwork to Delete:", style="Header.TLabel").pack(anchor="w")

        # Fetch artworks
        try:
            response = requests.get(
                f"{BASE_URL}/artworks/?format=json",
                headers=api_headers(),
                timeout=15
            )
            response.raise_for_status()
            artwork_list = response.json().get("artworks", [])
            artwork_titles = [art.get("title", "") for art in artwork_list]
        except Exception as e:
            messagebox.showerror("Error", f"Failed to fetch artworks: {str(e)}")
            delete_window.destroy()
            return

        sel_var = tk.StringVar()
        dropdown = ttk.Combobox(frm, textvariable=sel_var, values=artwork_titles, state="readonly")
        dropdown.pack(fill=tk.X, pady=(10, 6))

        info = ttk.Label(frm, text="", style="Subtle.TLabel")
        info.pack(anchor="w")

        def confirm_delete():
            title = sel_var.get()
            if not title:
                return
            try:
                r = requests.post(
                    f"{BASE_URL}/artworks/delete_artwork/",
                    json={"title": title},
                    headers=api_headers(),
                    timeout=20
                )
                if r.status_code == 200:
                    messagebox.showinfo("Success", f"“{title}” deleted successfully!")
                    info.config(text=f"Deleted “{title}”.")
                else:
                    messagebox.showerror("Error", f"Failed to delete artwork: {r.text}")
            except Exception as ex:
                messagebox.showerror("Error", f"Request failed: {ex}")

        btns = ttk.Frame(frm)
        btns.pack(fill=tk.X, pady=(12, 0))
        ttk.Button(btns, text="Delete from ArtSite", style="Danger.TButton", command=confirm_delete).pack(side=tk.LEFT)
        ttk.Button(btns, text="Close", command=delete_window.destroy).pack(side=tk.RIGHT)

    # -------------- Upload --------------
    def upload_artworks(self):
        if self.current_results is None or self.current_results.empty:
            messagebox.showwarning("Warning", "No artworks to upload. Do a search first!")
            return

        if len(self.current_results) > 4:
            messagebox.showwarning("Warning", "Can only upload a max of 4 items at a time. Narrow your search.")
            return

        for _, row in self.current_results.iterrows():
            title = row.get("Title", "")
            price = simpledialog.askstring("Enter Sale Price", f"Enter sale price for ‘{title}’:")
            if not price:
                continue

            images = []
            while True:
                image_path = filedialog.askopenfilename(
                    title=f"Select Image for ‘{title}’",
                    filetypes=[("Image Files", "*.JPG *.jpeg *.png")]
                )
                if image_path:
                    images.append(image_path)
                    more = messagebox.askyesno("Upload More?", "Upload another image for this artwork?")
                    if not more:
                        break
                else:
                    break

            if not images:
                continue

            # Detect sheet size flag "S" in Height/Width
            raw_height = str(row.get('Height', '')).strip()
            raw_width = str(row.get('Width', '')).strip()

            if raw_height.startswith('S'):
                height_clean = raw_height.replace('S', '').strip()
                width_clean = raw_width.replace('S', '').strip()
                dimensions_text = ""
                sheet_size = f"{height_clean}\" x {width_clean}\""
            else:
                dimensions_text = f"{raw_height}\" x {raw_width}\""
                sheet_size = ""

            data = {
                'artist': row.get('Artist', ''),
                'title': title,
                'year': row.get('Year', ''),
                'medium': row.get('Medium', ''),
                'description': row.get('Description/Notes', ''),
                'dimensions_text': dimensions_text,
                'sheet_size': sheet_size,
                'catalog_number': row.get('Catalog Number', ''),
                'price': price,
            }

            files_dict = {}
            try:
                for i, img in enumerate(images):
                    mime = "image/jpeg"
                    if img.lower().endswith(".png"):
                        mime = "image/png"
                    with open(img, 'rb') as f:
                        files_dict[f'image_{i}'] = (os.path.basename(img), f.read(), mime)

                response = requests.post(
                    f"{BASE_URL}/artworks/upload_artwork/",
                    data=data,
                    files=files_dict,
                    headers=api_headers(),
                    timeout=60
                )
                if response.status_code == 201:
                    messagebox.showinfo("Success", f"‘{title}’ uploaded successfully!")
                    self._set_status(f"Uploaded “{title}”.")
                    if not self.open_listings_button:
                        # Optional floating button is less needed now; use status & menu
                        pass
                else:
                    messagebox.showerror("Error", f"Upload failed: {response.text}")
            except Exception as e:
                messagebox.showerror("Error", f"Upload failed: {e}")

    def upload_sales_sheet(self):
        """
        Uploads the currently loaded sales spreadsheet (or asks you to choose one)
        to the server endpoint that imports it into the DB for Pieces Sold.
        """
        if not self.sales_path or not os.path.exists(self.sales_path):
            pick = filedialog.askopenfilename(
                title="Select Sales Excel Spreadsheet to Upload",
                filetypes=[("Excel files", "*.xlsx *.xls")]
            )
            if not pick:
                return
            self.sales_path = pick

        try:
            with open(self.sales_path, "rb") as f:
                files = {"file": (os.path.basename(self.sales_path), f,
                                  "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")}
                r = requests.post(
                    SALES_SHEET_ENDPOINT,
                    files=files,
                    headers=api_headers(),
                    timeout=120
                )

            if r.status_code in (200, 201):
                messagebox.showinfo("Success", "Sales spreadsheet uploaded and imported!")
                self._set_status("Sales sheet uploaded → Pieces Sold updated.")
            else:
                messagebox.showerror("Upload Error", f"Sales upload failed ({r.status_code}):\n{r.text}")

        except Exception as e:
            messagebox.showerror("Upload Error", f"Sales upload failed:\n{e}")

    # -------------- Main --------------


def main():
    root = tk.Tk()
    app = ArtCatalogApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
